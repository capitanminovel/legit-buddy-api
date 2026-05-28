"""Renders products.json into a fully static HTML file — no JS fetch needed."""
import json
from pathlib import Path
from datetime import datetime, timezone, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

CST          = timezone(timedelta(hours=-6))
DATA         = Path(__file__).parent / "docs" / "products.json"
STRAINS_DATA = Path(__file__).parent / "docs" / "strains_enriched.json"
OUT          = Path(__file__).parent / "docs" / "index.html"
NEW_DAYS     = 3
SOLD_DAYS    = 2

CAT_ICONS = {
    "flower":"🌿","pre-roll":"🚬","pre_roll":"🚬","preroll":"🚬",
    "edible":"🍬","edibles":"🍬","concentrate":"💎","concentrates":"💎",
    "vape":"💨","vapes":"💨","cartridge":"💨","cartridges":"💨",
    "tincture":"💧","tinctures":"💧","topical":"🧴","topicals":"🧴",
    "capsule":"💊","capsules":"💊","accessory":"🛠️","accessories":"🛠️",
    "beverage":"🥤","beverages":"🥤",
}
TIER_LABELS = {
    "gram":"1g","two_gram":"2g","eighth":"⅛ oz",
    "quarter":"¼ oz","half_ounce":"½ oz","ounce":"1 oz","unit":"Unit",
}

def cat_icon(c): return CAT_ICONS.get((c or "").lower().strip(), "🌱")

def strain_class(s):
    t = (s or "").lower()
    if "indica" in t: return "strain-indica"
    if "sativa" in t: return "strain-sativa"
    if "hybrid" in t: return "strain-hybrid"
    if "cbd"    in t: return "strain-cbd"
    if "cbg"    in t: return "strain-cbg"
    return "strain-default"

def age_days(iso):
    try:
        dt = datetime.fromisoformat(iso)
        return (datetime.now(CST) - dt).days
    except Exception:
        return 999

def new_badge(first_seen):
    d = age_days(first_seen)
    if d == 0:          return '<span class="new-badge">New Today</span>'
    if d <= NEW_DAYS:   return f'<span class="recent-badge">New ({d}d ago)</span>'
    return ""

def build_card(p, key):
    ci = cat_icon(p.get("category", ""))
    img = (f'<img src="{p["image"]}" alt="{p["name"]}" loading="lazy" onerror="this.parentNode.innerHTML=\'<div class=no-img>{ci}</div>\'">'
           if p.get("image")
           else f'<div class="no-img">{ci}</div>')

    strain_b = (f'<span class="strain-badge {strain_class(p["strain_type"])}">{p["strain_type"]}</span>'
                if p.get("strain_type") else "")
    age_b    = new_badge(p.get("first_seen", ""))
    badges   = f'<div class="badges">{age_b}{strain_b}</div>' if (strain_b or age_b) else ""

    thc_pill = f'<span class="potency-pill thc">THC {p["thc"]}</span>' if p.get("thc") else ""
    cbd_pill = f'<span class="potency-pill cbd">CBD {p["cbd"]}</span>' if p.get("cbd") else ""
    potency  = f'<div class="potency-row">{thc_pill}{cbd_pill}</div>' if (thc_pill or cbd_pill) else ""

    terps  = "".join(f'<span class="terp">{t}</span>' for t in (p.get("terpenes") or [])[:4])
    terp_h = f'<div class="terp-row">{terps}</div>' if terps else ""

    minors = " · ".join(filter(None, [
        f'CBG {p["cbg"]}' if p.get("cbg") else "",
        f'CBN {p["cbn"]}' if p.get("cbn") else "",
    ]))
    minor_h = f'<div style="font-size:.68rem;color:var(--muted);margin-top:2px">{minors}</div>' if minors else ""

    tiers = p.get("price_tiers") or {}
    if tiers:
        chips = "".join(f'<div class="tier">{v}<span>{TIER_LABELS.get(k,k)}</span></div>'
                        for k,v in tiers.items())
        price_h = f'<div class="price-tiers">{chips}</div>'
    elif p.get("price"):
        price_h = f'<div class="price-single">{p["price"]}</div>'
    else:
        price_h = ""

    weight_h = f'<div class="card-weight">{p["weight"]}</div>' if p.get("weight") else ""
    brand_h  = f'<div class="card-brand">{p["brand"]}</div>'   if p.get("brand")  else ""

    terpenes_csv = ",".join(p.get("terpenes") or [])

    return f"""
    <div class="card" data-key="{key}" data-terpenes="{terpenes_csv}" onclick="openModal('{key}')">
      <div class="card-img">{img}{badges}{potency}<div class="rating-badge"></div></div>
      <div class="card-body">
        {brand_h}
        <div class="card-name">{p["name"]}</div>
        {weight_h}{minor_h}{terp_h}
        <div class="price-section">{price_h}</div>
        <div class="card-detail-hint">Tap for strain guide →</div>
      </div>
    </div>"""

def build():
    with open(DATA) as f:
        db = json.load(f)

    strains = {}
    if STRAINS_DATA.exists():
        with open(STRAINS_DATA) as f:
            strains = json.load(f)

    schedule = {"shifts": [], "last_updated": None}
    sched_path = Path(__file__).parent / "docs" / "schedule.json"
    if sched_path.exists():
        with open(sched_path) as f:
            schedule = json.load(f)

    now     = datetime.now(CST)
    ts      = now.strftime("%a, %b %d %Y — %I:%M %p CST")
    TARGET  = ("flower", "pre-roll", "vapes", "edibles")
    all_p   = [(k,v) for k,v in db["products"].items()
               if v.get("in_stock", True) and (v.get("category","").lower() in TARGET)]

    all_p.sort(key=lambda x: (age_days(x[1].get("first_seen","")), x[1].get("name","")))

    from collections import defaultdict
    cats = defaultdict(list)
    for k, p in all_p:
        cats[p.get("category") or "Other"].append((k, p))

    new_items = [(k, p) for k, p in all_p if age_days(p.get("first_seen","")) <= NEW_DAYS]

    new_section = ""
    if new_items:
        new_cards = "".join(build_card(p, k) for k, p in new_items)
        n = len(new_items)
        new_section = f"""
    <section class="section new-arrivals-section" data-cat="all">
      <div class="new-arrivals-head">
        <span class="new-arrivals-title">✨ New in the Last 3 Days</span>
        <span class="new-arrivals-count" data-total="{n}">{n} product{"s" if n!=1 else ""}</span>
      </div>
      <div class="grid">{new_cards}</div>
    </section>
    <div class="section-divider" data-cat="all"></div>"""

    sold_items = sorted(
        [(k, p) for k, p in db["products"].items()
         if not p.get("in_stock", True)
         and p.get("category", "").lower() in TARGET
         and age_days(p.get("last_seen", "")) <= SOLD_DAYS],
        key=lambda x: age_days(x[1].get("last_seen", ""))
    )

    sold_section = ""
    if sold_items:
        def sold_row(p):
            d = age_days(p.get("last_seen", ""))
            when = "Today" if d == 0 else f"{d}d ago"
            ci   = cat_icon(p.get("category", ""))
            thc  = f'<span class="sold-thc">THC {p["thc"]}</span>' if p.get("thc") else ""
            sb   = f'<span class="strain-badge {strain_class(p["strain_type"])} sold-strain">{p["strain_type"]}</span>' if p.get("strain_type") else ""
            return (f'<div class="sold-row">'
                    f'<span class="sold-icon">{ci}</span>'
                    f'<span class="sold-name">{p["name"]}</span>'
                    f'{sb}{thc}'
                    f'<span class="sold-when">Gone {when}</span>'
                    f'</div>')
        rows = "".join(sold_row(p) for _, p in sold_items)
        n = len(sold_items)
        sold_section = f"""
    <section class="sold-section" data-cat="all">
      <div class="sold-head">
        <span class="sold-title">🚫 Sold Out — Last 2 Days</span>
        <span class="sold-count">{n} item{"s" if n!=1 else ""}</span>
      </div>
      <div class="sold-list">{rows}</div>
    </section>
    <div class="section-divider" data-cat="all"></div>"""

    tab_btns = '<button class="tab on" data-cat="all" onclick="filterCat(this)">All Products</button>\n'
    tab_btns += "\n".join(
        f'<button class="tab" data-cat="{c.lower()}" onclick="filterCat(this)">{cat_icon(c)} {c}</button>'
        for c in sorted(cats)
    )

    sections = ""
    for cat in sorted(cats):
        items = cats[cat]
        cards = "".join(build_card(p, k) for k, p in items)
        n = len(items)
        sections += f"""
    <section class="section" data-cat="{cat.lower()}">
      <div class="section-head">
        <span class="section-title">{cat_icon(cat)} {cat}</span>
        <span class="section-count" data-total="{n}">{n} product{"s" if n!=1 else ""}</span>
      </div>
      <div class="grid">{cards}</div>
    </section>"""

    # Embed all product data + strain enrichment + schedule as JS
    products_js  = json.dumps({k: v for k, v in db["products"].items()}, ensure_ascii=False)
    strains_js   = json.dumps(strains, ensure_ascii=False)
    schedule_js  = json.dumps(schedule, ensure_ascii=False)
    tab_btns    += '\n<button class="tab sched-tab" data-cat="schedule" onclick="openScheduleTab(this)">📅 Schedule</button>'

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>MN Legit Cannabis – South Metro Menu</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&family=Nunito:wght@700;800;900&family=Nunito+Sans:wght@400;600&display=swap" rel="stylesheet">
  <style>
    *,*::before,*::after{{box-sizing:border-box;margin:0;padding:0}}
    :root{{
      --brand:#1a7a4a;--brand-lt:#e8f5ee;--text:#111827;--muted:#6b7280;
      --border:#e5e7eb;--bg:#f3f4f6;--white:#ffffff;
      --indica:#7c3aed;--sativa:#d97706;--hybrid:#0891b2;--cbd:#2563eb;--cbg:#6366f1;
      --new:#16a34a;--radius:10px;
      --sg-green:#3d5c2e;--sg-pink:#e88fa2;--sg-cream:#f5f0e8;
      --sg-dark:#2a3f1f;--sg-border:#4a7030;
    }}
    body.dark{{
      --brand:#4ade80;--brand-lt:#14261e;--text:#e2e8f0;--muted:#94a3b8;
      --border:#2d3f32;--bg:#0d1a11;--white:#132019;
      --indica:#a78bfa;--sativa:#fbbf24;--hybrid:#38bdf8;--cbd:#60a5fa;--cbg:#818cf8;
      --new:#4ade80;
    }}
    body.dark header,body.dark .tabs-wrap,body.dark .legend,body.dark footer{{background:var(--white);border-color:var(--border)}}
    body.dark .card{{background:#1a2d20;border-color:var(--border)}}
    body.dark .mood-bar{{background:#1a2d20;border-color:var(--border)}}
    body.dark .mood-chip{{background:#132019;color:var(--text);border-color:var(--border)}}
    body.dark .mood-chip:hover{{background:#1e3826;border-color:var(--brand)}}
    body.dark .search-input{{background:#132019;color:var(--text);border-color:var(--border)}}
    body.dark .search-input:focus{{background:#1a2d20;border-color:var(--brand)}}
    body.dark .new-arrivals-section{{background:linear-gradient(135deg,#132019,#0d1f14);border-color:#2d4a35}}
    body.dark .sold-section{{background:linear-gradient(135deg,#1c1505,#1f1a08);border-color:#78350f}}
    body.dark .sold-row{{background:rgba(255,255,255,.05)}}
    body.dark .terp{{background:#0d2015;color:#4ade80;border-color:#1e4a2a}}
    body.dark .tier{{background:#132019;border-color:var(--border)}}
    body.dark .card.match-strong{{border-left:5px solid #4ade80;box-shadow:-2px 0 10px rgba(74,222,128,.3)}}
    body.dark .card.match-good{{border-left:5px solid #fbbf24;box-shadow:none}}
    body.dark .card.match-weak{{border-left:5px solid #475569;box-shadow:none}}
    body.dark .top-bar{{background:#0d2015}}
    body.dark .modal-box,.dark .profile-box{{background:#1a2d20}}
    body.dark .sg-card{{background:#1e3826;border-color:#3a6040}}
    body.dark .sg-name,.dark .sg-row strong{{color:#c8f5d4}}
    body.dark .sg-row{{color:#b2c9b8}}
    body.dark .profile-box{{background:#132019}}
    body.dark .profile-header{{background:#132019;border-color:#2d4a35}}
    body{{font-family:'Inter',system-ui,sans-serif;background:var(--bg);color:var(--text);min-height:100vh;font-size:14px}}
    .top-bar{{background:var(--brand);color:#fff;text-align:center;font-size:.75rem;padding:6px;letter-spacing:.3px}}
    header{{background:var(--white);border-bottom:1px solid var(--border);padding:0 24px;position:sticky;top:0;z-index:30}}
    .header-inner{{max-width:1400px;margin:0 auto;display:flex;align-items:center;gap:20px;height:64px}}
    .logo{{display:flex;align-items:center;gap:10px;font-weight:700;font-size:1.05rem;color:var(--brand);text-decoration:none;white-space:nowrap}}
    .logo-leaf{{width:34px;height:34px;background:var(--brand);border-radius:50% 50% 50% 0;display:flex;align-items:center;justify-content:center;font-size:1rem;color:#fff;flex-shrink:0}}
    .mascot-wrap{{flex-shrink:0;cursor:default;user-select:none;display:flex;align-items:center}}
    .mascot-wrap img{{height:52px;width:auto;display:block}}
    .mascot-flip{{transform:scaleX(-1)}}
    .dark-toggle{{margin-left:auto;background:none;border:1.5px solid var(--border);border-radius:20px;padding:5px 12px;font-size:.78rem;font-weight:600;cursor:pointer;color:var(--muted);font-family:inherit;transition:all .15s;white-space:nowrap;flex-shrink:0}}
    .dark-toggle:hover{{border-color:var(--brand);color:var(--brand)}}
    .header-meta{{margin-left:auto;text-align:right;font-size:.75rem;color:var(--muted);line-height:1.5}}
    .header-meta strong{{color:var(--brand)}}
    .tabs-wrap{{background:var(--white);border-bottom:1px solid var(--border);position:sticky;top:64px;z-index:20}}
    .tabs{{max-width:1400px;margin:0 auto;display:flex;gap:2px;overflow-x:auto;padding:0 24px;scrollbar-width:none}}
    .tabs::-webkit-scrollbar{{display:none}}
    .tab{{flex-shrink:0;padding:12px 16px;border:none;background:none;font-family:inherit;font-size:.82rem;font-weight:500;color:var(--muted);cursor:pointer;border-bottom:2px solid transparent;transition:color .15s,border-color .15s;white-space:nowrap}}
    .tab:hover{{color:var(--brand)}}
    .tab.on{{color:var(--brand);border-bottom-color:var(--brand)}}
    .legend{{display:flex;gap:14px;flex-wrap:wrap;align-items:center;padding:10px 24px;background:var(--white);border-bottom:1px solid var(--border);font-size:.74rem;color:var(--muted)}}
    .legend-item{{display:flex;align-items:center;gap:5px}}
    main{{max-width:1400px;margin:0 auto;padding:28px 24px 100px}}
    .section{{margin-bottom:44px}}
    .section-head{{display:flex;align-items:baseline;gap:10px;margin-bottom:18px}}
    .section-title{{font-size:1.1rem;font-weight:700}}
    .section-count{{font-size:.8rem;color:var(--muted)}}
    .grid{{display:grid;grid-template-columns:repeat(auto-fill,minmax(200px,1fr));gap:16px}}
    .card{{background:var(--white);border:1px solid var(--border);border-radius:var(--radius);overflow:hidden;display:flex;flex-direction:column;transition:box-shadow .15s,transform .15s;cursor:pointer}}
    .card:hover{{box-shadow:0 4px 16px rgba(0,0,0,.12);transform:translateY(-2px)}}
    .card:hover .card-detail-hint{{opacity:1}}
    .card-detail-hint{{font-size:.67rem;color:var(--brand);font-weight:600;text-align:center;padding:4px 0 0;opacity:0;transition:opacity .15s;letter-spacing:.2px}}
    .card-img{{position:relative;background:#f9fafb;border-bottom:1px solid var(--border);height:170px;overflow:hidden;display:flex;align-items:center;justify-content:center}}
    .card-img img{{width:100%;height:100%;object-fit:cover;display:block}}
    .no-img{{font-size:3.2rem;color:#d1d5db}}
    .badges{{position:absolute;top:8px;left:8px;display:flex;flex-direction:column;gap:4px}}
    .strain-badge{{display:inline-block;padding:2px 8px;border-radius:20px;font-size:.66rem;font-weight:700;letter-spacing:.3px;text-transform:uppercase;color:#fff}}
    .strain-indica{{background:var(--indica)}}.strain-sativa{{background:var(--sativa)}}
    .strain-hybrid{{background:var(--hybrid)}}.strain-cbd{{background:var(--cbd)}}
    .strain-cbg{{background:var(--cbg)}}.strain-default{{background:#6b7280}}
    .new-badge{{display:inline-block;padding:2px 8px;border-radius:20px;font-size:.66rem;font-weight:700;letter-spacing:.3px;text-transform:uppercase;background:var(--new);color:#fff}}
    .recent-badge{{display:inline-block;padding:2px 8px;border-radius:20px;font-size:.66rem;font-weight:700;letter-spacing:.3px;text-transform:uppercase;background:#f59e0b;color:#fff}}
    .potency-row{{position:absolute;bottom:8px;right:8px;display:flex;gap:4px}}
    .potency-pill{{background:rgba(0,0,0,.65);color:#fff;font-size:.67rem;font-weight:600;padding:2px 6px;border-radius:4px}}
    .potency-pill.thc{{background:rgba(22,163,74,.85)}}.potency-pill.cbd{{background:rgba(37,99,235,.85)}}
    .card-body{{padding:12px 13px 14px;flex:1;display:flex;flex-direction:column;gap:4px}}
    .card-brand{{font-size:.72rem;color:var(--muted);font-weight:500;text-transform:uppercase;letter-spacing:.3px}}
    .card-name{{font-size:.9rem;font-weight:600;line-height:1.3;color:var(--text)}}
    .card-weight{{font-size:.74rem;color:var(--muted)}}
    .terp-row{{display:flex;gap:4px;flex-wrap:wrap;margin-top:3px}}
    .terp{{font-size:.65rem;background:#f0fdf4;color:var(--brand);border:1px solid #bbf7d0;padding:1px 6px;border-radius:10px}}
    .effects-row{{display:flex;gap:4px;flex-wrap:wrap;margin-top:2px}}
    .effect{{font-size:.65rem;background:#eff6ff;color:#3b82f6;border:1px solid #bfdbfe;padding:1px 6px;border-radius:10px}}
    .price-section{{margin-top:auto;padding-top:10px}}
    .price-single{{font-size:1rem;font-weight:700;color:var(--brand)}}
    .price-tiers{{display:flex;gap:5px;flex-wrap:wrap}}
    .tier{{font-size:.7rem;font-weight:500;border:1px solid var(--border);border-radius:5px;padding:3px 7px;color:var(--text);background:#fafafa}}
    .tier span{{display:block;font-size:.62rem;color:var(--muted)}}
    footer{{text-align:center;padding:16px 20px 80px;font-size:.72rem;color:var(--muted);border-top:1px solid var(--border);background:var(--white)}}
    .footer-sticky{{position:fixed;bottom:0;left:0;right:0;background:var(--white);border-top:1px solid var(--border);padding:8px 16px;display:flex;justify-content:space-between;align-items:center;font-size:.72rem;color:var(--muted);z-index:200;box-shadow:0 -2px 8px rgba(0,0,0,.06)}}
    .footer-sticky .fs-stock{{font-weight:700;color:var(--brand)}}
    .footer-sticky .fs-updated{{font-size:.68rem}}
    @media(min-width:768px){{.footer-sticky{{display:none}}}}
    @media(max-width:600px){{
      .header-meta{{display:none}}
      .mascot-wrap img{{height:36px}}
      .dark-toggle{{padding:4px 8px;font-size:.7rem}}
      .header-inner{{gap:10px}}
    }}
    .new-arrivals-section{{background:linear-gradient(135deg,#f0fdf4,#ecfdf5);border:2px solid #86efac;border-radius:12px;padding:20px;margin-bottom:32px}}
    .new-arrivals-head{{display:flex;align-items:baseline;gap:10px;margin-bottom:18px}}
    .new-arrivals-title{{font-size:1.15rem;font-weight:700;color:var(--new)}}
    .new-arrivals-count{{font-size:.8rem;color:var(--muted)}}
    .sold-section{{background:linear-gradient(135deg,#fffbeb,#fef3c7);border:2px solid #fcd34d;border-radius:12px;padding:20px;margin-bottom:32px}}
    .sold-head{{display:flex;align-items:baseline;gap:10px;margin-bottom:14px}}
    .sold-title{{font-size:1.15rem;font-weight:700;color:#b45309}}
    .sold-count{{font-size:.8rem;color:var(--muted)}}
    .sold-list{{display:flex;flex-direction:column;gap:8px}}
    .sold-row{{display:flex;align-items:center;gap:8px;padding:8px 10px;background:rgba(255,255,255,.6);border-radius:8px;flex-wrap:wrap}}
    .sold-icon{{font-size:1.1rem;flex-shrink:0}}
    .sold-name{{font-weight:600;font-size:.9rem;flex:1;min-width:140px}}
    .sold-strain{{font-size:.65rem;padding:2px 6px}}
    .sold-thc{{font-size:.75rem;color:#92400e;font-weight:600;background:#fde68a;padding:2px 7px;border-radius:10px}}
    .sold-when{{margin-left:auto;font-size:.75rem;color:#b45309;font-weight:700;white-space:nowrap}}
    .section-divider{{height:2px;background:linear-gradient(90deg,var(--brand-lt),transparent);margin:0 0 36px;border-radius:1px}}
    .hidden{{display:none!important}}

    /* ── Mood / Effect filter bar ── */
    .mood-bar{{background:var(--white);border:1px solid var(--border);border-radius:12px;padding:14px 18px;margin-bottom:24px;display:flex;flex-direction:column;gap:10px}}
    .mood-bar-top{{display:flex;align-items:center;gap:10px;flex-wrap:wrap}}
    .mood-bar-label{{font-size:.78rem;font-weight:700;color:var(--muted);text-transform:uppercase;letter-spacing:.4px;white-space:nowrap}}
    .mood-chips{{display:flex;gap:7px;flex-wrap:wrap;flex:1}}
    .mood-chip{{border:1.5px solid var(--border);background:var(--bg);color:var(--text);border-radius:20px;padding:6px 13px;font-size:.78rem;font-weight:600;cursor:pointer;transition:all .15s;white-space:nowrap;font-family:inherit}}
    .mood-chip:hover{{border-color:var(--sg-green);color:var(--sg-green);background:#f0fdf4}}
    .mood-chip.on{{background:var(--sg-green);color:var(--sg-pink);border-color:var(--sg-green);font-weight:700}}
    .mood-clear{{border:none;background:none;color:var(--muted);font-size:.78rem;font-weight:600;cursor:pointer;padding:6px 8px;border-radius:20px;white-space:nowrap;font-family:inherit}}
    .mood-clear:hover{{color:#e53e3e}}
    .mood-status{{font-size:.78rem;color:var(--sg-green);font-weight:500;padding:2px 0 0;line-height:1.45}}
    .mood-status strong{{font-weight:700}}
    .mood-zero{{font-size:.85rem;color:var(--muted);text-align:center;padding:32px 0;font-weight:500}}
    .card.match-strong{{border-left:5px solid #16a34a;box-shadow:-2px 0 10px rgba(22,163,74,.2)}}
    .card.match-good{{border-left:5px solid #d97706}}
    .card.match-weak{{border-left:5px solid #94a3b8}}

    /* ── Rating badge shown on cards when mood active ── */
    .rating-badge{{display:none;position:absolute;top:8px;right:8px;min-width:28px;height:28px;border-radius:50%;font-family:'Nunito',sans-serif;font-weight:900;font-size:13px;align-items:center;justify-content:center;z-index:2;box-shadow:0 2px 6px rgba(0,0,0,.25);border:2px solid rgba(255,255,255,.6)}}
    .rating-badge.show{{display:flex}}
    .rating-badge.rb-strong{{background:#16a34a;color:#fff}}
    .rating-badge.rb-good{{background:#d97706;color:#fff}}
    .rating-badge.rb-weak{{background:#94a3b8;color:#fff}}
    body.dark .card.match-strong{{border-left:5px solid #4ade80;box-shadow:-2px 0 10px rgba(74,222,128,.3)}}
    body.dark .card.match-good{{border-left:5px solid #fbbf24}}
    body.dark .card.match-weak{{border-left:5px solid #475569}}
    body.dark .rating-badge.rb-strong{{background:#4ade80;color:#0d1a11}}
    body.dark .rating-badge.rb-good{{background:#fbbf24;color:#1a1000}}
    body.dark .rating-badge.rb-weak{{background:#475569;color:#e2e8f0}}

    /* ── Moods info button + modal ── */
    .mood-info-btn{{border:none;background:none;color:var(--muted);font-size:.8rem;font-weight:600;cursor:pointer;padding:4px 6px;border-radius:8px;font-family:inherit;white-space:nowrap;flex-shrink:0}}
    .mood-info-btn:hover{{color:var(--brand);background:var(--brand-lt)}}
    .moods-modal-overlay{{position:fixed;inset:0;background:rgba(0,0,0,.55);z-index:500;display:flex;align-items:flex-end;justify-content:center}}
    .moods-modal-overlay.hidden{{display:none}}
    .moods-modal-box{{background:var(--bg);width:100%;max-width:680px;max-height:88vh;border-radius:20px 20px 0 0;overflow-y:auto;padding:0 0 30px}}
    .moods-modal-head{{position:sticky;top:0;background:var(--bg);border-bottom:1px solid var(--border);padding:16px 22px;display:flex;align-items:center;justify-content:space-between}}
    .moods-modal-title{{font-family:'Nunito',sans-serif;font-weight:900;font-size:1.1rem;color:var(--brand)}}
    .moods-modal-close{{background:none;border:none;font-size:1.4rem;cursor:pointer;color:var(--muted);line-height:1}}
    .mood-card{{background:var(--white);border:1px solid var(--border);border-radius:12px;padding:14px 18px;margin:16px 18px 0}}
    .mood-card-head{{display:flex;align-items:center;gap:10px;margin-bottom:6px}}
    .mood-card-icon{{font-size:1.4rem}}
    .mood-card-name{{font-family:'Nunito',sans-serif;font-weight:900;font-size:1rem;color:var(--brand)}}
    .mood-card-science{{font-size:.78rem;color:var(--text);line-height:1.55;margin-bottom:8px}}
    .mood-card-terps{{display:flex;gap:5px;flex-wrap:wrap}}
    .mood-card-terp{{font-size:.68rem;font-weight:700;background:var(--brand-lt);color:var(--brand);border:1px solid #bbf7d0;border-radius:10px;padding:2px 8px}}
    body.dark .mood-card{{background:#1a2d20;border-color:var(--border)}}
    body.dark .moods-modal-box{{background:var(--bg)}}
    body.dark .moods-modal-head{{background:var(--bg)}}

    /* ── Text search ── */
    .search-row{{display:flex;align-items:center;gap:8px}}
    .search-wrap{{position:relative;flex:1;max-width:520px}}
    .search-icon{{position:absolute;left:11px;top:50%;transform:translateY(-50%);font-size:.85rem;pointer-events:none}}
    .search-input{{width:100%;border:1.5px solid var(--border);border-radius:24px;padding:8px 36px 8px 32px;font-size:.82rem;font-family:inherit;outline:none;background:var(--bg);color:var(--text);transition:border-color .15s}}
    .search-input:focus{{border-color:var(--sg-green);background:#fff}}
    .search-input::placeholder{{color:#aaa}}
    .search-clear{{position:absolute;right:10px;top:50%;transform:translateY(-50%);border:none;background:none;color:#aaa;font-size:1rem;cursor:pointer;padding:2px;line-height:1}}
    .search-clear:hover{{color:#e53e3e}}

    @media(max-width:640px){{
      header{{padding:0 14px}}.tabs{{padding:0 14px}}main{{padding:18px 14px 100px}}.legend{{padding:10px 14px}}
      .grid{{grid-template-columns:repeat(auto-fill,minmax(155px,1fr));gap:12px}}.card-img{{height:140px}}
      .mood-bar{{padding:12px 14px}}.mood-chips{{gap:5px}}.mood-chip{{font-size:.73rem;padding:5px 10px}}
    }}

    /* ── Modal overlay ── */
    .modal-overlay{{position:fixed;inset:0;background:rgba(0,0,0,.55);z-index:1000;display:flex;align-items:center;justify-content:center;padding:16px;opacity:0;pointer-events:none;transition:opacity .2s}}
    .modal-overlay.open{{opacity:1;pointer-events:all}}
    .modal-box{{background:#e8e0d0;border-radius:18px;max-width:620px;width:100%;max-height:90vh;overflow-y:auto;position:relative;transform:scale(.95);transition:transform .2s;font-family:'Nunito Sans',sans-serif}}
    .modal-overlay.open .modal-box{{transform:scale(1)}}
    .modal-close{{position:sticky;top:12px;float:right;margin:12px 16px 0 0;background:var(--sg-green);color:var(--sg-pink);border:none;border-radius:50%;width:32px;height:32px;font-size:1.1rem;cursor:pointer;font-weight:900;line-height:32px;text-align:center;z-index:10;flex-shrink:0}}
    .modal-close:hover{{background:var(--sg-dark)}}
    .modal-inner{{padding:16px 22px 22px;clear:both}}

    /* ── Strain card (inside modal) — matches legit_strain_guide.html ── */
    .sg-card{{background:white;border:3px solid var(--sg-border);border-radius:16px;padding:18px 22px;margin-bottom:14px}}
    .sg-name{{font-family:'Nunito',sans-serif;font-weight:900;font-size:22px;text-align:center;text-transform:uppercase;letter-spacing:.05em;color:var(--sg-dark);margin-bottom:2px}}
    .sg-type{{text-align:center;font-size:12.5px;font-weight:700;color:#555;margin-bottom:4px}}
    .sg-supplier{{display:block;background:var(--sg-green);color:var(--sg-pink);font-family:'Nunito',sans-serif;font-weight:800;font-size:10.5px;letter-spacing:.07em;text-transform:uppercase;border-radius:20px;padding:3px 10px;width:fit-content;margin:0 auto 10px}}
    .sg-divider{{border:none;border-top:2px solid var(--sg-border);margin:8px 0 12px}}
    .sg-row{{font-size:12.5px;line-height:1.55;margin-bottom:4px;color:#222}}
    .sg-row strong{{font-weight:700;color:var(--sg-dark);font-family:'Nunito',sans-serif;font-size:12.5px}}
    .sg-thc-cbd{{display:flex;gap:8px;justify-content:center;margin-bottom:8px;flex-wrap:wrap}}
    .sg-pill{{font-size:11px;font-weight:700;padding:2px 10px;border-radius:20px;font-family:'Nunito',sans-serif}}
    .sg-pill.thc{{background:#16a34a;color:#fff}}.sg-pill.cbd{{background:#2563eb;color:#fff}}
    .sg-price{{text-align:center;font-size:13px;font-weight:700;color:var(--sg-dark);margin-bottom:6px}}
    .modal-actions{{display:flex;gap:10px;margin-top:16px;justify-content:center;flex-wrap:wrap}}
    .btn-add-profile{{background:var(--sg-green);color:var(--sg-pink);border:none;border-radius:24px;padding:10px 22px;font-family:'Nunito',sans-serif;font-weight:800;font-size:13px;letter-spacing:.05em;text-transform:uppercase;cursor:pointer;transition:background .15s}}
    .btn-add-profile:hover{{background:var(--sg-dark)}}
    .btn-add-profile.added{{background:#6b7280;color:#fff}}
    .btn-close-modal{{background:transparent;color:var(--sg-green);border:2px solid var(--sg-green);border-radius:24px;padding:10px 22px;font-family:'Nunito',sans-serif;font-weight:800;font-size:13px;letter-spacing:.05em;text-transform:uppercase;cursor:pointer}}

    /* ── Profile floating button ── */
    .profile-fab{{position:fixed;bottom:28px;right:24px;background:var(--sg-green);color:var(--sg-pink);border:none;border-radius:30px;padding:12px 20px;font-family:'Nunito',sans-serif;font-weight:900;font-size:13px;letter-spacing:.05em;text-transform:uppercase;cursor:pointer;box-shadow:0 4px 20px rgba(0,0,0,.25);z-index:500;display:none;align-items:center;gap:8px;transition:background .15s,transform .1s}}
    .profile-fab:hover{{background:var(--sg-dark);transform:scale(1.04)}}
    .profile-fab-count{{background:var(--sg-pink);color:var(--sg-dark);border-radius:50%;width:22px;height:22px;display:inline-flex;align-items:center;justify-content:center;font-size:11px;font-weight:900}}

    /* ── Profile drawer ── */
    .profile-drawer{{position:fixed;inset:0;background:rgba(0,0,0,.55);z-index:1100;display:none;align-items:flex-end;justify-content:center}}
    .profile-drawer.open{{display:flex}}
    .profile-box{{background:#e8e0d0;width:100%;max-width:680px;max-height:85vh;border-radius:18px 18px 0 0;overflow-y:auto;padding:0 0 30px}}
    .profile-header{{display:flex;align-items:center;justify-content:space-between;padding:18px 22px 14px;background:#e8e0d0;position:sticky;top:0;border-bottom:2px solid var(--sg-border)}}
    .profile-header-title{{font-family:'Nunito',sans-serif;font-weight:900;font-size:18px;color:var(--sg-dark);text-transform:uppercase;letter-spacing:.05em}}
    .profile-header-actions{{display:flex;gap:8px;flex-wrap:wrap}}
    .btn-export{{background:var(--sg-green);color:var(--sg-pink);border:none;border-radius:20px;padding:8px 16px;font-family:'Nunito',sans-serif;font-weight:800;font-size:11.5px;letter-spacing:.05em;text-transform:uppercase;cursor:pointer}}
    .btn-export:hover{{background:var(--sg-dark)}}
    .export-bar{{display:flex;gap:8px;flex-wrap:wrap;align-items:center;padding:8px 16px;background:var(--white);border-bottom:1px solid var(--border)}}
    .export-bar-label{{font-size:.7rem;color:var(--muted);font-weight:600;text-transform:uppercase;letter-spacing:.04em;white-space:nowrap}}
    .btn-export-all{{background:#1a7a4a;color:#fff;border:none;border-radius:20px;padding:6px 14px;font-family:'Nunito',sans-serif;font-weight:800;font-size:11px;letter-spacing:.04em;text-transform:uppercase;cursor:pointer;white-space:nowrap}}
    .btn-export-all:hover{{background:#145e38}}
    .btn-export-avail{{background:var(--sg-pink);color:#fff;border:none;border-radius:20px;padding:6px 14px;font-family:'Nunito',sans-serif;font-weight:800;font-size:11px;letter-spacing:.04em;text-transform:uppercase;cursor:pointer;white-space:nowrap}}
    .btn-export-avail:hover{{background:#d4708a}}
    body.dark .export-bar{{background:var(--white);border-color:var(--border)}}
    body.dark .btn-export-all{{background:#4ade80;color:#0d1a11}}
    body.dark .btn-export-all:hover{{background:#22c55e}}
    body.dark .btn-export-avail{{background:#e88fa2;color:#1a0a0e}}
    .btn-staff-guide{{background:transparent;color:var(--muted);border:1.5px solid var(--border);border-radius:20px;padding:6px 14px;font-family:'Nunito',sans-serif;font-weight:800;font-size:11px;letter-spacing:.04em;text-transform:uppercase;cursor:pointer;white-space:nowrap}}
    .btn-staff-guide:hover{{border-color:var(--brand);color:var(--brand)}}

    /* ── Staff guide modal ── */
    .sg-guide-overlay{{position:fixed;inset:0;background:rgba(0,0,0,.55);z-index:500;display:flex;align-items:flex-end;justify-content:center}}
    .sg-guide-overlay.hidden{{display:none}}
    .sg-guide-box{{background:var(--bg);width:100%;max-width:720px;max-height:90vh;border-radius:20px 20px 0 0;overflow-y:auto;padding:0 0 40px}}
    .sg-guide-head{{position:sticky;top:0;background:var(--bg);border-bottom:1px solid var(--border);padding:16px 22px;display:flex;align-items:center;justify-content:space-between;z-index:2}}
    .sg-guide-title{{font-family:'Nunito',sans-serif;font-weight:900;font-size:1.1rem;color:var(--brand)}}
    .sg-guide-close{{background:none;border:none;font-size:1.4rem;cursor:pointer;color:var(--muted);line-height:1}}
    .sg-guide-section{{margin:20px 18px 0}}
    .sg-guide-section-title{{font-family:'Nunito',sans-serif;font-weight:900;font-size:.9rem;color:var(--brand);text-transform:uppercase;letter-spacing:.06em;border-bottom:2px solid var(--brand-lt);padding-bottom:6px;margin-bottom:12px}}
    .sg-guide-card{{background:var(--white);border:1px solid var(--border);border-radius:12px;padding:14px 16px;margin-bottom:10px}}
    .sg-guide-card-head{{display:flex;align-items:center;gap:8px;margin-bottom:5px}}
    .sg-guide-card-icon{{font-size:1.2rem}}
    .sg-guide-card-name{{font-family:'Nunito',sans-serif;font-weight:900;font-size:.95rem;color:var(--text)}}
    .sg-guide-card-body{{font-size:.8rem;color:var(--text);line-height:1.6}}
    .sg-guide-card-body strong{{color:var(--brand)}}
    .sg-guide-card-body em{{color:var(--muted);font-style:normal;font-size:.75rem}}
    .sg-guide-tag{{display:inline-block;font-size:.68rem;font-weight:700;background:var(--brand-lt);color:var(--brand);border:1px solid #bbf7d0;border-radius:10px;padding:2px 8px;margin:2px 2px 0 0}}
    .sg-guide-table{{width:100%;border-collapse:collapse;font-size:.78rem;margin-top:6px}}
    .sg-guide-table th{{text-align:left;color:var(--muted);font-weight:700;padding:4px 8px 4px 0;border-bottom:1px solid var(--border)}}
    .sg-guide-table td{{padding:6px 8px 6px 0;border-bottom:1px solid var(--border);vertical-align:top;line-height:1.45}}
    .sg-guide-table td:first-child{{font-weight:700;white-space:nowrap;color:var(--brand)}}
    .sg-guide-note{{background:#fffbeb;border:1px solid #fde68a;border-radius:10px;padding:10px 14px;font-size:.78rem;color:#92400e;line-height:1.55;margin-top:10px}}
    body.dark .sg-guide-box{{background:var(--bg)}}
    body.dark .sg-guide-head{{background:var(--bg)}}
    body.dark .sg-guide-card{{background:#1a2d20;border-color:var(--border)}}
    body.dark .sg-guide-note{{background:#1a1500;border-color:#713f12;color:#fde68a}}
    .export-popup-overlay{{position:fixed;inset:0;background:rgba(0,0,0,.6);z-index:9999;display:flex;align-items:center;justify-content:center}}
    .export-popup-overlay.hidden{{display:none}}
    .export-popup-box{{background:#e8e0d0;border:3px solid #4a7030;border-radius:20px;padding:28px 32px;text-align:center;max-width:320px;width:90%;box-shadow:0 20px 60px rgba(0,0,0,.4);animation:epPop .35s cubic-bezier(.175,.885,.32,1.275)}}
    @keyframes epPop{{from{{transform:scale(.7);opacity:0}}to{{transform:scale(1);opacity:1}}}}
    .export-popup-gif{{width:190px;height:190px;object-fit:cover;border-radius:14px;border:3px solid #4a7030;margin-bottom:14px}}
    .export-popup-title{{font-family:'Nunito',sans-serif;font-weight:900;font-size:20px;color:#2a3f1f;text-transform:uppercase;letter-spacing:.05em;margin-bottom:4px}}
    .export-popup-sub{{font-size:13px;color:#3d5c2e;font-weight:600;margin-bottom:18px}}
    .export-popup-btn{{background:#3d5c2e;color:#e88fa2;border:none;border-radius:20px;padding:10px 28px;font-family:'Nunito',sans-serif;font-weight:900;font-size:13px;letter-spacing:.05em;text-transform:uppercase;cursor:pointer;transition:transform .1s}}
    .export-popup-btn:hover{{background:#2a3f1f}}
    .export-popup-btn.ready{{animation:epPulse 0.9s ease-in-out infinite;background:#4a7030}}
    @keyframes epPulse{{0%,100%{{transform:scale(1)}}50%{{transform:scale(1.06)}}}}
    .btn-clear{{background:transparent;color:#888;border:1px solid #ccc;border-radius:20px;padding:8px 14px;font-family:'Nunito',sans-serif;font-weight:700;font-size:11px;text-transform:uppercase;cursor:pointer}}
    .btn-close-drawer{{background:transparent;color:var(--sg-green);border:2px solid var(--sg-green);border-radius:20px;padding:8px 14px;font-family:'Nunito',sans-serif;font-weight:800;font-size:11px;text-transform:uppercase;cursor:pointer}}
    .profile-cards{{padding:16px 18px 0}}
    .profile-empty{{text-align:center;padding:40px;color:#888;font-family:'Nunito',sans-serif;font-weight:700;font-size:14px}}
    .profile-item-remove{{position:absolute;top:10px;right:12px;background:transparent;border:none;color:#aaa;font-size:18px;cursor:pointer;font-weight:700;line-height:1}}
    .profile-item-remove:hover{{color:#e53e3e}}
    @media(max-width:640px){{
      .modal-box{{max-height:95vh;border-radius:14px}}
      .profile-box{{max-height:92vh}}
      .modal-inner{{padding:12px 14px 18px}}
    }}
    /* ── Schedule ── */
    .sched-tab{{margin-left:auto}}
    #scheduleSection{{display:none;padding:24px}}
    #scheduleSection.active{{display:block}}
    .sched-img-section{{margin-bottom:20px;border-bottom:1px solid var(--border);padding-bottom:16px}}
    .sched-img-toggle{{background:none;border:1px solid var(--border);border-radius:20px;padding:7px 16px;font-size:.8rem;font-weight:600;color:var(--muted);cursor:pointer;transition:all .15s}}
    .sched-img-toggle:hover{{border-color:var(--brand);color:var(--brand)}}
    .sched-img-wrap{{margin-top:16px}}
    .sched-img-label{{font-weight:700;font-size:.85rem;color:var(--muted);margin:16px 0 8px}}
    .sched-img{{width:100%;border-radius:10px;border:1px solid var(--border);display:block}}
    .pin-overlay{{position:fixed;inset:0;background:rgba(0,0,0,.55);backdrop-filter:blur(4px);z-index:300;display:flex;align-items:center;justify-content:center}}
    .pin-overlay.hidden{{display:none}}
    .pin-box{{background:var(--white);border-radius:16px;padding:36px 40px;text-align:center;min-width:280px;box-shadow:0 20px 60px rgba(0,0,0,.25)}}
    .pin-box h2{{font-family:'Nunito',sans-serif;font-size:1.25rem;color:var(--text);margin-bottom:6px}}
    .pin-box p{{font-size:.8rem;color:var(--muted);margin-bottom:20px}}
    .pin-input{{width:120px;text-align:center;font-size:1.8rem;letter-spacing:.4em;font-weight:700;padding:8px 12px;border:2px solid var(--border);border-radius:10px;font-family:'Nunito',sans-serif;color:var(--text);outline:none}}
    .pin-input:focus{{border-color:var(--brand)}}
    .pin-error{{color:#dc2626;font-size:.78rem;margin-top:10px;min-height:18px}}
    .sched-header{{display:flex;align-items:center;gap:12px;margin-bottom:20px;flex-wrap:wrap}}
    .sched-title{{font-family:'Nunito',sans-serif;font-size:1.2rem;font-weight:800;color:var(--text)}}
    .sched-nav{{display:flex;align-items:center;gap:8px;margin-left:auto}}
    .sched-nav-btn{{padding:6px 14px;border:1px solid var(--border);border-radius:20px;background:var(--white);color:var(--muted);cursor:pointer;font-size:.78rem;font-weight:600;transition:all .15s}}
    .sched-nav-btn:hover{{border-color:var(--brand);color:var(--brand)}}
    .sched-nav-btn.disabled{{opacity:.4;cursor:default;pointer-events:none}}
    .sched-week-label{{font-size:.82rem;color:var(--muted);font-weight:600;min-width:90px;text-align:center}}
    .sched-filters{{display:flex;gap:6px;flex-wrap:wrap;margin-bottom:18px}}
    .sched-filter-btn{{padding:5px 14px;border-radius:20px;border:1px solid var(--border);background:var(--white);color:var(--muted);cursor:pointer;font-size:.75rem;font-weight:600;transition:all .15s}}
    .sched-filter-btn:hover{{border-color:var(--brand);color:var(--brand)}}
    .sched-filter-btn.on{{background:var(--brand-lt);border-color:var(--brand);color:var(--brand)}}
    .sched-day{{margin-bottom:18px;background:var(--white);border:1px solid var(--border);border-radius:12px;overflow:hidden}}
    .sched-day.today .sched-day-head{{background:var(--brand);color:#fff}}
    .sched-day-head{{padding:9px 16px;font-size:.82rem;font-weight:700;background:var(--bg);color:var(--text);display:flex;align-items:center;gap:8px}}
    .sched-today-badge{{background:#fff;color:var(--brand);font-size:.65rem;font-weight:800;padding:2px 8px;border-radius:10px;text-transform:uppercase;letter-spacing:.05em}}
    .sched-shifts{{padding:4px 0}}
    .sched-shift{{display:flex;align-items:center;gap:12px;padding:9px 16px;border-bottom:1px solid var(--bg);font-size:.82rem}}
    .sched-shift:last-child{{border-bottom:none}}
    .sched-shift-name{{font-weight:600;color:var(--text);min-width:130px}}
    .sched-shift-time{{color:var(--muted);white-space:nowrap}}
    .sched-empty{{text-align:center;padding:48px 24px;color:var(--muted);font-size:.88rem}}
    .sched-empty-icon{{font-size:2rem;margin-bottom:12px;opacity:.4}}
    .sched-updated{{font-size:.7rem;color:var(--muted);text-align:right;margin-top:12px}}
  </style>
</head>
<body>
<div class="top-bar">🌿 MN Legit Cannabis · South Metro · Updated daily at 4:30 PM CST</div>
<header>
  <div class="header-inner">
    <div class="mascot-wrap">
      <img src="https://media1.giphy.com/media/v1.Y2lkPTc5MGI3NjExb216bXVxcmVpMm1zNTR3NWxob3hoeXloYWFlbWswcW13NnZ3MnlsMCZlcD12MV9pbnRlcm5hbF9naWZfYnlfaWQmY3Q9cw/5C29m5sRSYch78Ko3B/giphy.gif" alt="" loading="lazy">
    </div>
    <a class="logo" href="#">
      <div><div>Legit Cannabis</div><div style="font-size:.7rem;font-weight:400;color:var(--muted)">South Metro</div></div>
    </a>
    <div class="mascot-wrap">
      <img src="https://media2.giphy.com/media/v1.Y2lkPTc5MGI3NjExOHAxcXpnaHdqN3QyaGtmaDM5azkyd2hlZndvNHZocncwamd6Y2Z5aSZlcD12MV9pbnRlcm5hbF9naWZfYnlfaWQmY3Q9Zw/VRvFAP4CXxUQw/giphy.gif" alt="" loading="lazy">
    </div>
    <button class="dark-toggle" id="darkToggle" onclick="toggleDark()">🌙 Dark</button>
    <div class="header-meta">
      <div>Last updated: <strong>{ts}</strong></div>
      <div>{len(all_p)} products in stock</div>
    </div>
  </div>
</header>
<div class="export-bar">
  <span class="export-bar-label">⬇ Export .docx:</span>
  <button class="btn-export-avail" onclick="showExportPopup('legit-available-guide.docx')">✅ Available Now ({len(all_p)} products)</button>
  <button class="btn-export-all"   onclick="showExportPopup('legit-master-guide.docx')">📦 Master Cache (all strains)</button>
  <button class="btn-staff-guide"  onclick="openStaffGuide()" style="margin-left:auto">📖 Staff Guide</button>
</div>

<!-- Export popup -->
<div class="export-popup-overlay hidden" id="exportPopup">
  <div class="export-popup-box">
    <img class="export-popup-gif" id="exportPopupGif" src="https://media.giphy.com/media/VK2JbAI71xTxlSVNNu/giphy.gif" alt="">
    <div class="export-popup-title">Your guide is ready 🍃</div>
    <div class="export-popup-sub" id="exportPopupSub">Downloading in <span id="exportCountdown">4</span>s…</div>
    <button class="export-popup-btn" id="exportGoBtn">Let's Go ⬇</button>
  </div>
</div>
<div class="legend">
  <div class="legend-item"><span class="strain-badge strain-indica">Indica</span></div>
  <div class="legend-item"><span class="strain-badge strain-sativa">Sativa</span></div>
  <div class="legend-item"><span class="strain-badge strain-hybrid">Hybrid</span></div>
  <div class="legend-item"><span class="strain-badge strain-cbd">CBD</span></div>
  <div class="legend-item"><span class="new-badge">New Today</span> Added today</div>
  <div class="legend-item"><span class="recent-badge">New (2d)</span> Within 3 days</div>
  <div class="legend-item"><span class="sold-thc" style="background:#fde68a;color:#92400e;padding:2px 7px;border-radius:10px;font-size:.75rem;font-weight:600">🚫 Gone</span> Sold out in last 2 days</div>
  <div class="legend-item" style="margin-left:auto;color:var(--brand);font-weight:600">Tap any product for strain guide →</div>
</div>
<div class="tabs-wrap"><div class="tabs" id="tabs">{tab_btns}</div></div>
<main>
  <div class="mood-bar" id="moodBar">
    <div class="mood-bar-top">
      <span class="mood-bar-label">Find your vibe</span>
      <div class="mood-chips" id="moodChips">
        <button class="mood-chip" data-mood="wind-down"      onclick="filterMood(this)" title="Myrcene + Linalool — muscle relaxation, sedation, GABAergic calm (Russo 2011)">😴 Wind Down</button>
        <button class="mood-chip" data-mood="anxiety-relief" onclick="filterMood(this)" title="Caryophyllene + Linalool + Limonene — the Kamal anxiolytic chemotype (Kamal et al. 2018, Front Neurosci)">🧘 Anxiety Relief</button>
        <button class="mood-chip" data-mood="lift-up"        onclick="filterMood(this)" title="Limonene + Terpinolene + Ocimene — mood elevation, citrus-forward uplift">⬆ Lift Up</button>
        <button class="mood-chip" data-mood="get-creative"   onclick="filterMood(this)" title="Pinene + Terpinolene — AChE inhibition sharpens focus; terpinolene drives cerebral creativity">🎨 Get Creative</button>
        <button class="mood-chip" data-mood="get-social"     onclick="filterMood(this)" title="Limonene + Terpinolene — euphoria, giggles, sociability without heavy sedation">😄 Get Social</button>
        <button class="mood-chip" data-mood="pain-body"      onclick="filterMood(this)" title="Caryophyllene (CB2 agonist) + Myrcene + Humulene — anti-inflammatory, analgesic, muscle relaxant">💆 Pain &amp; Body</button>
        <button class="mood-chip" data-mood="just-happy"     onclick="filterMood(this)" title="Limonene + Linalool — balanced euphoria and body warmth">✨ Just Happy</button>
        <button class="mood-chip" data-mood="aphrodisiac"    onclick="filterMood(this)" title="Limonene (dopamine↑) + Linalool (anxiety↓) + Geraniol (rose terpene) + Caryophyllene (CB2 tactile) — 3,000 years of documented use (Russo 2011)">🌹 Aphrodisiac</button>
      </div>
      <button class="mood-clear hidden" id="moodClear" onclick="clearMood()">✕ Mood</button>
      <button class="mood-info-btn" onclick="openMoodsInfo()" title="Learn the science behind each mood filter">ℹ️ How it works</button>
    </div>
    <div class="search-row">
      <div class="search-wrap">
        <span class="search-icon">🔍</span>
        <input type="text" id="searchInput" class="search-input"
               placeholder="Search: anxiety, PTSD, sleep, pain, creative, Myrcene…"
               oninput="handleSearch(this.value)">
        <button class="search-clear hidden" id="searchClear" onclick="clearSearch()">✕</button>
      </div>
    </div>
    <div class="mood-status hidden" id="moodStatus"></div>
  </div>
  {new_section}{sold_section}{sections}
  <div class="mood-zero hidden" id="moodZero">No products match this vibe right now — try another filter.</div>
</main>
<footer>
  Last updated: {ts} &nbsp;·&nbsp; {len(all_p)} products in stock &nbsp;·&nbsp; MN Legit Cannabis South Metro
</footer>
<div class="footer-sticky">
  <span class="fs-stock">{len(all_p)} products in stock</span>
  <span class="fs-updated">Updated {ts}</span>
</div>

<!-- Moods info modal -->
<div class="moods-modal-overlay hidden" id="moodsInfoModal" onclick="if(event.target===this)closeMoodsInfo()">
  <div class="moods-modal-box">
    <div class="moods-modal-head">
      <span class="moods-modal-title">🔬 How Moods Are Scored</span>
      <button class="moods-modal-close" onclick="closeMoodsInfo()">✕</button>
    </div>
    <div style="padding:10px 18px 0;font-size:.8rem;color:var(--muted);line-height:1.55">
      Every rating is derived from <strong>COA terpene lab data only</strong> — not dispensary marketing copy.
      Scores 1–10 are generated by Claude AI using published pharmacology research (Russo 2011 <em>Br J Pharmacol</em>, Kamal 2018 <em>Front Neurosci</em>, Gertsch 2008 <em>PNAS</em>).
      Cards sort best → weakest match. Border color = strength at a glance.
    </div>
    <div id="moodsInfoCards"></div>
  </div>
</div>

<!-- Staff guide modal -->
<div class="sg-guide-overlay hidden" id="staffGuideModal" onclick="if(event.target===this)closeStaffGuide()">
  <div class="sg-guide-box">
    <div class="sg-guide-head">
      <span class="sg-guide-title">📖 Staff Guide</span>
      <button class="sg-guide-close" onclick="closeStaffGuide()">✕</button>
    </div>
    <div id="staffGuideContent"></div>
  </div>
</div>

<!-- Strain modal -->
<div class="modal-overlay" id="strainModal" onclick="closeModalOutside(event)">
  <div class="modal-box">
    <button class="modal-close" onclick="closeModal()">✕</button>
    <div class="modal-inner">
      <div id="modalCard"></div>
      <div class="modal-actions">
        <button class="btn-add-profile" id="btnAddProfile" onclick="toggleProfile()">＋ Add to Profile</button>
        <button class="btn-close-modal" onclick="closeModal()">Close</button>
      </div>
    </div>
  </div>
</div>

<!-- Profile drawer -->
<div class="profile-drawer" id="profileDrawer">
  <div class="profile-box">
    <div class="profile-header">
      <div class="profile-header-title">📋 Strain Profile</div>
      <div class="profile-header-actions">
        <button class="btn-export" onclick="exportGuide()">⬇ Download Strain Guide</button>
        <button class="btn-clear" onclick="clearProfile()">Clear All</button>
        <button class="btn-close-drawer" onclick="closeDrawer()">Close</button>
      </div>
    </div>
    <div class="profile-cards" id="profileCards">
      <div class="profile-empty" id="profileEmpty">No strains added yet.<br>Tap any product card, then "Add to Profile."</div>
    </div>
  </div>
</div>

<!-- Floating profile button -->
<button class="profile-fab" id="profileFab" onclick="openDrawer()">
  📋 My Profile <span class="profile-fab-count" id="fabCount">0</span>
</button>

<script>
const PRODUCTS = {products_js};
const STRAINS  = {strains_js};

// ── Mood map: effects + terpenes that predict each vibe ──
// Sources: Russo 2011 Br J Pharmacol; Kamal et al. 2018 Front Neurosci;
//          Smith et al. 2022 PLOS ONE (terpenes > indica/sativa label)
const MOOD_MAP = {{
  'wind-down': {{
    label: 'Wind Down',
    science: 'Myrcene + Linalool stack — sedation, muscle relaxation, GABAergic calm',
    effects:  ['Sleepy','Relaxing','Calming','Chill','Body High','Unbothered'],
    terpenes: ['Myrcene','Linalool']
  }},
  'anxiety-relief': {{
    label: 'Anxiety Relief',
    science: 'Caryophyllene (CB2) + Linalool (GABA) + Limonene (5-HT1A) — Kamal 2018 anxiolytic chemotype',
    effects:  ['Calming','Chill','Relaxing','Unbothered','Blissful'],
    terpenes: ['Caryophyllene','Linalool','Limonene']
  }},
  'lift-up': {{
    label: 'Lift Up',
    science: 'Limonene mood elevation (Komori 1995) + Terpinolene cerebral uplift + Ocimene/Valencene citrus energy',
    effects:  ['Uplifting','Euphoric','Happy','Blissful','Energetic'],
    terpenes: ['Limonene','Terpinolene','Ocimene','Valencene']
  }},
  'get-creative': {{
    label: 'Get Creative',
    science: 'Pinene AChE inhibition sharpens memory + Terpinolene cerebral drive (Miyazawa & Yamafuji 2005)',
    effects:  ['Creative','Cerebral','Focused'],
    terpenes: ['Pinene','B Pinene','Terpinolene','Limonene']
  }},
  'get-social': {{
    label: 'Get Social',
    science: 'Limonene + Terpinolene — euphoria and giggles without heavy sedation',
    effects:  ['Social','Giggly','Talkative','Happy','Euphoric'],
    terpenes: ['Limonene','Terpinolene']
  }},
  'pain-body': {{
    label: 'Pain & Body',
    science: 'Caryophyllene (CB2 agonist, Gertsch 2008 PNAS) + Myrcene analgesic + Humulene anti-inflammatory',
    effects:  ['Body High','Tingly','Relaxing'],
    terpenes: ['Caryophyllene','Myrcene','Humulene','Linalool','Bisabolol']
  }},
  'just-happy': {{
    label: 'Just Happy',
    science: 'Limonene + Linalool + Terpinolene — balanced euphoria and body warmth',
    effects:  ['Happy','Euphoric','Blissful','Giggly','Tingly','Uplifting'],
    terpenes: ['Limonene','Linalool','Terpinolene']
  }},
  'aphrodisiac': {{
    label: 'Aphrodisiac',
    science: 'Limonene (dopamine/serotonin ↑) + Linalool (anxiety ↓, the #1 arousal blocker) + Geraniol (rose terpene, historic aphrodisiac) + Caryophyllene (CB2 tactile sensitivity) + Terpinolene (lowers inhibitions). Cannabis aphrodisiac use documented across cultures for 3,000+ years — Russo 2011.',
    effects:  ['Aroused','Tingly','Euphoric','Blissful','Happy'],
    terpenes: ['Limonene','Linalool','Geraniol','Terpinolene','Ocimene','Caryophyllene']
  }}
}};

let currentKey   = null;
let profileKeys  = [];
let activeMood   = null;
let activeCat    = 'all';
let activeSearch = '';
let searchTimer  = null;

// Research-backed terpene → effect map (Russo 2011, Kamal 2018, Smith 2022)
// Terpenes come from COA data — the only source we fully trust.
// Scraped "effects" from the dispensary page are NOT used anywhere.
const TERPENE_EFFECTS = {{
  'Myrcene':       ['Relaxing','Sleepy','Body High','Calming','Hungry'],
  'Limonene':      ['Uplifting','Happy','Euphoric','Energetic','Focused','Aroused'],
  'Caryophyllene': ['Calming','Relaxing','Body High','Tingly','Aroused'],
  'Linalool':      ['Sleepy','Calming','Relaxing','Chill','Blissful','Aroused'],
  'Pinene':        ['Focused','Creative','Energetic','Uplifting','Cerebral'],
  'B Pinene':      ['Focused','Creative','Energetic','Uplifting','Cerebral'],
  'Terpinolene':   ['Creative','Uplifting','Euphoric','Energetic','Cerebral','Giggly','Aroused'],
  'Humulene':      ['Calming','Body High'],
  'Ocimene':       ['Uplifting','Energetic','Social','Aroused'],
  'Valencene':     ['Uplifting','Happy','Social'],
  'Bisabolol':     ['Calming','Relaxing','Blissful'],
  'Geraniol':      ['Calming','Happy','Blissful','Aroused','Tingly'],
  'Terpinene':     ['Uplifting','Energetic'],
}};

function derivedEffects(terpenes) {{
  const set = new Set();
  terpenes.forEach(t => (TERPENE_EFFECTS[t] || []).forEach(e => set.add(e)));
  return [...set];
}}

// Terpene position → concentration proxy (earlier in COA list = dominant)
function terpenePositionScore(tx, moodTerpenes) {{
  return moodTerpenes.reduce((sum, mt) => {{
    const idx = tx.indexOf(mt);
    if (idx === -1) return sum;
    if (idx === 0)  return sum + 4.0;   // dominant terpene — very strong signal
    if (idx <= 1)   return sum + 2.5;   // secondary
    if (idx <= 3)   return sum + 1.5;   // tertiary
    return sum + 0.75;                  // minor trace
  }}, 0);
}}

function moodScore(card, mood) {{
  if (!mood) return 0;
  const tx      = (card.dataset.terpenes || '').split(',').filter(Boolean);
  const derived = derivedEffects(tx);
  const terpScore   = terpenePositionScore(tx, mood.terpenes);
  const effectScore = mood.effects.filter(e => derived.includes(e)).length * 0.8;
  // Scale to 0-10 with realistic spread: max raw ~8-9 → caps at 10
  return Math.min(10, Math.round((terpScore + effectScore) * 10 / 9));
}}

// Save original DOM order on load so we can restore it
document.addEventListener('DOMContentLoaded', () => {{
  document.querySelectorAll('.grid').forEach(grid => {{
    [...grid.children].forEach((el, i) => {{ el.dataset.origIndex = i; }});
  }});
}});

function fmtList(arr) {{
  if (!arr || !arr.length) return '—';
  return arr.join(', ');
}}

function buildSgCard(key, forExport) {{
  const p = PRODUCTS[key] || {{}};
  const s = STRAINS[key]  || {{}};
  const thcPill = p.thc ? `<span class="sg-pill thc">THC ${{p.thc}}</span>` : '';
  const cbdPill = p.cbd ? `<span class="sg-pill cbd">CBD ${{p.cbd}}</span>` : '';
  const pills   = (thcPill || cbdPill) ? `<div class="sg-thc-cbd">${{thcPill}}${{cbdPill}}</div>` : '';
  const price   = p.price ? `<div class="sg-price">${{p.price}}${{p.weight ? ' · ' + p.weight : ''}}</div>` : '';
  const removeBtn = forExport ? '' : `<button class="profile-item-remove" onclick="removeFromProfile('${{key}}')" title="Remove">✕</button>`;

  const tx      = p.terpenes || [];
  const derived = derivedEffects(tx);
  const rows = [
    s.lineage    ? `<div class="sg-row"><strong>Lineage:</strong> ${{s.lineage}}</div>` : '',
    derived.length ? `<div class="sg-row"><strong>Effects</strong> <span style="font-size:10px;color:#888;font-weight:400">(from COA terpenes)</span><strong>:</strong> ${{fmtList(derived)}}</div>` : '',
    p.flavors?.length ? `<div class="sg-row"><strong>Flavors:</strong> ${{fmtList(p.flavors)}}</div>` : '',
    tx.length    ? `<div class="sg-row"><strong>Terpenes:</strong> ${{fmtList(tx)}}</div>` : '',
    s.therapeutic ? `<div class="sg-row"><strong>Therapeutic:</strong> ${{s.therapeutic}}</div>` : '',
    s.negative   ? `<div class="sg-row"><strong>Negative:</strong> ${{s.negative}}</div>` : '',
    s.aroma      ? `<div class="sg-row"><strong>Aroma:</strong> ${{s.aroma}}</div>` : '',
    s.misc       ? `<div class="sg-row"><strong>Misc:</strong> ${{s.misc}}</div>` : '',
  ].join('');

  return `
  <div class="sg-card" style="position:relative">
    ${{removeBtn}}
    <div class="sg-name">${{p.name || 'Unknown'}}</div>
    <div class="sg-type">${{p.strain_type ? '— ' + p.strain_type : ''}}</div>
    <span class="sg-supplier">${{p.brand || 'Unknown'}}</span>
    ${{pills}}${{price}}
    <hr class="sg-divider">
    ${{rows}}
  </div>`;
}}

function openModal(key) {{
  currentKey = key;
  const p = PRODUCTS[key] || {{}};
  document.getElementById('modalCard').innerHTML = buildSgCard(key, false);
  const btn = document.getElementById('btnAddProfile');
  const inProfile = profileKeys.includes(key);
  btn.textContent = inProfile ? '✓ In Profile' : '＋ Add to Profile';
  btn.classList.toggle('added', inProfile);
  document.getElementById('strainModal').classList.add('open');
  document.body.style.overflow = 'hidden';
}}

function closeModal() {{
  document.getElementById('strainModal').classList.remove('open');
  document.body.style.overflow = '';
}}

function closeModalOutside(e) {{
  if (e.target === document.getElementById('strainModal')) closeModal();
}}

function toggleProfile() {{
  if (!currentKey) return;
  const idx = profileKeys.indexOf(currentKey);
  if (idx === -1) {{
    profileKeys.push(currentKey);
  }} else {{
    profileKeys.splice(idx, 1);
  }}
  const btn = document.getElementById('btnAddProfile');
  const inProfile = profileKeys.includes(currentKey);
  btn.textContent = inProfile ? '✓ In Profile' : '＋ Add to Profile';
  btn.classList.toggle('added', inProfile);
  updateFab();
}}

function removeFromProfile(key) {{
  profileKeys = profileKeys.filter(k => k !== key);
  updateFab();
  renderProfileCards();
}}

function updateFab() {{
  const fab = document.getElementById('profileFab');
  document.getElementById('fabCount').textContent = profileKeys.length;
  fab.style.display = profileKeys.length > 0 ? 'flex' : 'none';
}}

function renderProfileCards() {{
  const el = document.getElementById('profileCards');
  const empty = document.getElementById('profileEmpty');
  if (profileKeys.length === 0) {{
    empty.style.display = 'block';
    el.innerHTML = '';
    el.appendChild(empty);
    return;
  }}
  empty.style.display = 'none';
  el.innerHTML = profileKeys.map(k => buildSgCard(k, false)).join('');
}}

function openDrawer() {{
  renderProfileCards();
  document.getElementById('profileDrawer').classList.add('open');
  document.body.style.overflow = 'hidden';
}}

function closeDrawer() {{
  document.getElementById('profileDrawer').classList.remove('open');
  document.body.style.overflow = '';
}}

function clearProfile() {{
  profileKeys = [];
  updateFab();
  renderProfileCards();
}}

const EXPORT_POPUP_CSS = `
  .welcome-overlay{{position:fixed;inset:0;background:rgba(0,0,0,.55);z-index:9999;display:flex;align-items:center;justify-content:center;animation:fadeIn .3s ease}}
  @keyframes fadeIn{{from{{opacity:0}}to{{opacity:1}}}}
  .welcome-box{{background:#e8e0d0;border:3px solid #4a7030;border-radius:20px;padding:28px 32px;text-align:center;max-width:340px;width:90%;position:relative;box-shadow:0 20px 60px rgba(0,0,0,.35);animation:popIn .35s cubic-bezier(.175,.885,.32,1.275)}}
  @keyframes popIn{{from{{transform:scale(.7);opacity:0}}to{{transform:scale(1);opacity:1}}}}
  .welcome-gif{{width:200px;height:200px;object-fit:cover;border-radius:14px;margin-bottom:14px;border:3px solid #4a7030}}
  .welcome-title{{font-family:'Nunito',sans-serif;font-weight:900;font-size:20px;color:#2a3f1f;text-transform:uppercase;letter-spacing:.05em;margin-bottom:4px}}
  .welcome-sub{{font-size:13px;color:#3d5c2e;font-weight:600;margin-bottom:18px}}
  .welcome-close{{background:#3d5c2e;color:#e88fa2;border:none;border-radius:20px;padding:9px 24px;font-family:'Nunito',sans-serif;font-weight:900;font-size:13px;letter-spacing:.05em;text-transform:uppercase;cursor:pointer}}
  .welcome-close:hover{{background:#2a3f1f}}
`;
const EXPORT_POPUP_HTML = `
<div class="welcome-overlay" id="welcomeOverlay" onclick="if(event.target===this)dismissWelcome()">
  <div class="welcome-box">
    <img class="welcome-gif" src="https://media.giphy.com/media/VK2JbAI71xTxlSVNNu/giphy.gif" alt="Welcome">
    <div class="welcome-title">Welcome to Legit 🍃</div>
    <div class="welcome-sub">Your strain guide is ready. Enjoy!</div>
    <button class="welcome-close" onclick="dismissWelcome()">Let's Go</button>
  </div>
</div>
<script>
  function dismissWelcome(){{document.getElementById('welcomeOverlay').remove();}}
  setTimeout(dismissWelcome, 5000);
<\/script>
`;

function exportGuide() {{
  if (profileKeys.length === 0) return;
  const cards = profileKeys.map(k => buildSgCard(k, true)).join('');
  const today = new Date().toLocaleDateString('en-US', {{month:'long', day:'numeric', year:'numeric'}});
  const html = `<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Legit Cannabis – Strain Guide</title>
<link href="https://fonts.googleapis.com/css2?family=Nunito:wght@700;800;900&family=Nunito+Sans:wght@400;600&display=swap" rel="stylesheet">
<style>
  :root{{--green:#3d5c2e;--pink:#e88fa2;--cream:#f5f0e8;--dark-green:#2a3f1f;--border-green:#4a7030;--text:#1a1a1a}}
  *{{box-sizing:border-box;margin:0;padding:0}}
  body{{background:#e8e0d0;font-family:'Nunito Sans',sans-serif;padding:24px 16px;color:var(--text)}}
  .page{{max-width:720px;margin:0 auto}}
  .header{{display:flex;align-items:center;gap:16px;margin-bottom:28px}}
  .logo-badge{{background:var(--green);border-radius:14px;padding:10px 16px;display:flex;align-items:center;gap:8px}}
  .logo-badge .leaf{{font-size:20px}}
  .logo-badge .name{{font-family:'Nunito',sans-serif;font-weight:900;font-size:15px;color:var(--pink);line-height:1.1;letter-spacing:.02em;text-transform:uppercase}}
  .header-title{{font-family:'Nunito',sans-serif;font-weight:900;font-size:26px;color:var(--dark-green);letter-spacing:.04em;text-transform:uppercase}}
  .header-sub{{font-size:12px;color:var(--green);font-weight:600;letter-spacing:.06em;text-transform:uppercase;margin-top:2px}}
  .sg-card{{background:white;border:3px solid var(--border-green);border-radius:16px;padding:18px 22px;margin-bottom:18px}}
  .sg-name{{font-family:'Nunito',sans-serif;font-weight:900;font-size:22px;text-align:center;text-transform:uppercase;letter-spacing:.05em;color:var(--dark-green);margin-bottom:2px}}
  .sg-type{{text-align:center;font-size:12.5px;font-weight:700;color:#555;margin-bottom:4px}}
  .sg-supplier{{display:block;background:var(--green);color:var(--pink);font-family:'Nunito',sans-serif;font-weight:800;font-size:10.5px;letter-spacing:.07em;text-transform:uppercase;border-radius:20px;padding:3px 10px;width:fit-content;margin:0 auto 10px}}
  .sg-thc-cbd{{display:flex;gap:8px;justify-content:center;margin-bottom:8px;flex-wrap:wrap}}
  .sg-pill{{font-size:11px;font-weight:700;padding:2px 10px;border-radius:20px;font-family:'Nunito',sans-serif}}
  .sg-pill.thc{{background:#16a34a;color:#fff}}.sg-pill.cbd{{background:#2563eb;color:#fff}}
  .sg-price{{text-align:center;font-size:13px;font-weight:700;color:var(--dark-green);margin-bottom:6px}}
  .sg-divider{{border:none;border-top:2px solid var(--border-green);margin:8px 0 12px}}
  .sg-row{{font-size:12.5px;line-height:1.55;margin-bottom:4px;color:#222}}
  .sg-row strong{{font-weight:700;color:var(--dark-green);font-family:'Nunito',sans-serif;font-size:12.5px}}
  .profile-item-remove{{display:none}}
  .sg-row span{{font-size:10px;color:#888;font-weight:400}}
  @media print{{body{{background:white;padding:0}}.sg-card{{break-inside:avoid}}}}
  ${{EXPORT_POPUP_CSS}}
</style>
</head>
<body>
${{EXPORT_POPUP_HTML}}
<div class="page">
  <div class="header">
    <div class="logo-badge"><span class="leaf">🍃</span><div class="name">LEGIT<br>CANNABIS</div></div>
    <div><div class="header-title">Strain Guide</div><div class="header-sub">Staff Reference · ${{today}}</div></div>
  </div>
  ${{cards}}
</div>
</body>
</html>`;

  const blob = new Blob([html], {{type: 'text/html;charset=utf-8'}});
  const url  = URL.createObjectURL(blob);
  const a    = document.createElement('a');
  a.href     = url;
  a.download = 'legit-strain-guide.html';
  a.click();
  URL.revokeObjectURL(url);
}}

function exportAll(mode) {{
  // mode: 'avail' = only products in PRODUCTS (in-stock scraped set)
  //       'master' = all keys in STRAINS cache
  const today = new Date().toLocaleDateString('en-US', {{month:'long', day:'numeric', year:'numeric'}});

  let keys;
  if (mode === 'avail') {{
    keys = Object.keys(PRODUCTS);
  }} else {{
    // master: all enriched keys, fall back to PRODUCTS data for any missing
    keys = [...new Set([...Object.keys(STRAINS), ...Object.keys(PRODUCTS)])];
  }}

  // Sort: flower → pre-roll → vapes → other, then alpha within each
  const catOrder = ['flower','pre-roll','vapes','edibles'];
  keys.sort((a, b) => {{
    const pa = PRODUCTS[a] || {{}};
    const pb = PRODUCTS[b] || {{}};
    const ca = (pa.category || '').toLowerCase();
    const cb = (pb.category || '').toLowerCase();
    const ia = catOrder.indexOf(ca); const ib = catOrder.indexOf(cb);
    const oa = ia === -1 ? 99 : ia;  const ob = ib === -1 ? 99 : ib;
    if (oa !== ob) return oa - ob;
    return (pa.name || a).localeCompare(pb.name || b);
  }});

  const cards = keys.map(k => buildSgCard(k, true)).join('');
  const title = mode === 'avail' ? 'Available Now' : 'Master Strain Cache';
  const fname = mode === 'avail' ? 'legit-available-guide.html' : 'legit-master-guide.html';

  const html = `<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Legit Cannabis – ${{title}}</title>
<link href="https://fonts.googleapis.com/css2?family=Nunito:wght@700;800;900&family=Nunito+Sans:wght@400;600&display=swap" rel="stylesheet">
<style>
  :root{{--green:#3d5c2e;--pink:#e88fa2;--cream:#f5f0e8;--dark-green:#2a3f1f;--border-green:#4a7030;--text:#1a1a1a}}
  *{{box-sizing:border-box;margin:0;padding:0}}
  body{{background:#e8e0d0;font-family:'Nunito Sans',sans-serif;padding:24px 16px;color:var(--text)}}
  .page{{max-width:720px;margin:0 auto}}
  .header{{display:flex;align-items:center;gap:16px;margin-bottom:28px}}
  .logo-badge{{background:var(--green);border-radius:14px;padding:10px 16px;display:flex;align-items:center;gap:8px}}
  .logo-badge .leaf{{font-size:20px}}
  .logo-badge .name{{font-family:'Nunito',sans-serif;font-weight:900;font-size:15px;color:var(--pink);line-height:1.1;letter-spacing:.02em;text-transform:uppercase}}
  .header-title{{font-family:'Nunito',sans-serif;font-weight:900;font-size:26px;color:var(--dark-green);letter-spacing:.04em;text-transform:uppercase}}
  .header-sub{{font-size:12px;color:var(--green);font-weight:600;letter-spacing:.06em;text-transform:uppercase;margin-top:2px}}
  .section-heading{{font-family:'Nunito',sans-serif;font-weight:900;font-size:16px;color:var(--dark-green);text-transform:uppercase;letter-spacing:.06em;border-bottom:2px solid var(--border-green);padding-bottom:6px;margin:24px 0 14px}}
  .sg-card{{background:white;border:3px solid var(--border-green);border-radius:16px;padding:18px 22px;margin-bottom:18px}}
  .sg-name{{font-family:'Nunito',sans-serif;font-weight:900;font-size:22px;text-align:center;text-transform:uppercase;letter-spacing:.05em;color:var(--dark-green);margin-bottom:2px}}
  .sg-type{{text-align:center;font-size:12.5px;font-weight:700;color:#555;margin-bottom:4px}}
  .sg-supplier{{display:block;background:var(--green);color:var(--pink);font-family:'Nunito',sans-serif;font-weight:800;font-size:10.5px;letter-spacing:.07em;text-transform:uppercase;border-radius:20px;padding:3px 10px;width:fit-content;margin:0 auto 10px}}
  .sg-thc-cbd{{display:flex;gap:8px;justify-content:center;margin-bottom:8px;flex-wrap:wrap}}
  .sg-pill{{font-size:11px;font-weight:700;padding:2px 10px;border-radius:20px;font-family:'Nunito',sans-serif}}
  .sg-pill.thc{{background:#16a34a;color:#fff}}.sg-pill.cbd{{background:#2563eb;color:#fff}}
  .sg-price{{text-align:center;font-size:13px;font-weight:700;color:var(--dark-green);margin-bottom:6px}}
  .sg-divider{{border:none;border-top:2px solid var(--border-green);margin:8px 0 12px}}
  .sg-row{{font-size:12.5px;line-height:1.55;margin-bottom:4px;color:#222}}
  .sg-row strong{{font-weight:700;color:var(--dark-green);font-family:'Nunito',sans-serif;font-size:12.5px}}
  .profile-item-remove{{display:none}}
  @media print{{body{{background:white;padding:0}}.sg-card{{break-inside:avoid}}}}
  ${{EXPORT_POPUP_CSS}}
</style>
</head>
<body>
${{EXPORT_POPUP_HTML}}
<div class="page">
  <div class="header">
    <div class="logo-badge"><span class="leaf">🍃</span><div class="name">LEGIT<br>CANNABIS</div></div>
    <div><div class="header-title">${{title}}</div><div class="header-sub">Staff Reference · ${{today}}</div></div>
  </div>
  ${{cards}}
</div>
</body>
</html>`;

  const blob = new Blob([html], {{type: 'text/html;charset=utf-8'}});
  const url  = URL.createObjectURL(blob);
  const a    = document.createElement('a');
  a.href = url; a.download = fname; a.click();
  URL.revokeObjectURL(url);
}}

// ── Category + mood + text search combined filter ──
function applyFilters() {{
  const mood = activeMood ? MOOD_MAP[activeMood] : null;
  const q    = activeSearch;
  let totalVisible = 0;

  document.querySelectorAll('.card').forEach(card => {{
    const key = card.dataset.key;

    // Category check
    const section = card.closest('.section');
    const catOk = activeCat === 'all' || (section && section.dataset.cat === activeCat);

    // Mood check — driven entirely by COA terpenes, not dispensary effect labels
    let moodOk = true;
    if (mood) {{
      const tx      = (card.dataset.terpenes || '').split(',').filter(Boolean);
      const derived = derivedEffects(tx);
      moodOk = mood.effects.some(e  => derived.includes(e))
             || mood.terpenes.some(t => tx.includes(t));
    }}

    // Text search — terpenes (COA), derived effects (research), and enriched
    // strain fields (therapeutic, aroma, misc, lineage).
    // Scraped dispensary "effects" intentionally excluded — not trusted.
    let searchOk = true;
    if (q) {{
      const p = PRODUCTS[key] || {{}};
      const s = STRAINS[key]  || {{}};
      const tx = (p.terpenes || []);
      const blob = [
        p.name, p.brand, p.strain_type,
        tx.join(' '),
        derivedEffects(tx).join(' '),
        (p.flavors || []).join(' '),
        s.lineage, s.therapeutic, s.negative, s.aroma, s.misc
      ].filter(Boolean).join(' ').toLowerCase();
      searchOk = blob.includes(q);
    }}

    const visible = catOk && moodOk && searchOk;
    card.classList.toggle('hidden', !visible);

    // Rating badge + match border
    card.classList.remove('match-strong','match-good','match-weak');
    const rb = card.querySelector('.rating-badge');
    if (visible && mood) {{
      // Prefer Claude AI rating (0-10), fall back to computed score scaled to 0-10
      const moodKey = activeMood.replace(/-/g,'_');
      const claudeRating = STRAINS[key]?.mood_ratings?.[moodKey];
      const score10 = claudeRating != null
        ? claudeRating
        : Math.min(10, Math.round(moodScore(card, mood) * 1.5));

      if (rb) {{
        rb.textContent = score10;
        rb.className = 'rating-badge show ' +
          (score10 >= 7 ? 'rb-strong' : score10 >= 4 ? 'rb-good' : 'rb-weak');
      }}
      if      (score10 >= 7) card.classList.add('match-strong');
      else if (score10 >= 4) card.classList.add('match-good');
      else                   card.classList.add('match-weak');
    }} else {{
      if (rb) rb.className = 'rating-badge';
    }}

    if (visible) totalVisible++;
  }});

  // Sort each grid: mood active → best score first; no mood → restore original order
  document.querySelectorAll('.grid').forEach(grid => {{
    const cards = [...grid.querySelectorAll('.card')];
    if (mood) {{
      const moodKey = activeMood.replace(/-/g,'_');
      cards.sort((a, b) => {{
        const ra = STRAINS[a.dataset.key]?.mood_ratings?.[moodKey] ?? Math.min(10, moodScore(a, mood) * 1.5);
        const rb = STRAINS[b.dataset.key]?.mood_ratings?.[moodKey] ?? Math.min(10, moodScore(b, mood) * 1.5);
        const diff = rb - ra;
        return diff !== 0 ? diff : (parseInt(a.dataset.origIndex)||0) - (parseInt(b.dataset.origIndex)||0);
      }});
    }} else {{
      cards.sort((a, b) => (parseInt(a.dataset.origIndex)||0) - (parseInt(b.dataset.origIndex)||0));
    }}
    cards.forEach(c => grid.appendChild(c));
  }});

  // Update section visibility + counts
  document.querySelectorAll('.section').forEach(s => {{
    const catOk = activeCat === 'all' || s.dataset.cat === activeCat;
    if (!catOk) {{ s.classList.add('hidden'); return; }}
    const vis = s.querySelectorAll('.card:not(.hidden)').length;
    s.classList.toggle('hidden', vis === 0);
    const countEl = s.querySelector('[data-total]');
    if (countEl) {{
      const total = countEl.dataset.total;
      countEl.textContent = mood
        ? `${{vis}} / ${{total}} matching`
        : `${{total}} product${{total == 1 ? '' : 's'}}`;
    }}
  }});

  // Divider + new-arrivals + sold-out sections (all-only)
  document.querySelectorAll('.section-divider').forEach(d => {{
    d.classList.toggle('hidden', activeCat !== 'all');
  }});
  const soldSec = document.querySelector('.sold-section');
  if (soldSec) soldSec.classList.toggle('hidden', activeCat !== 'all');

  // Show "no results" message
  document.getElementById('moodZero').classList.toggle('hidden', totalVisible > 0);

  // Update status bar
  const statusEl = document.getElementById('moodStatus');
  const parts = [];
  if (mood)  parts.push(`<strong>${{mood.label}}:</strong> ${{mood.science}}`);
  if (q)     parts.push(`searching <strong>"${{q}}"</strong>`);
  if (parts.length) {{
    parts.push(`— ${{totalVisible}} product${{totalVisible == 1 ? '' : 's'}} found`);
    statusEl.innerHTML = parts.join(' · ');
    statusEl.classList.remove('hidden');
  }} else {{
    statusEl.classList.add('hidden');
  }}
}}

function handleSearch(val) {{
  clearTimeout(searchTimer);
  searchTimer = setTimeout(() => {{
    activeSearch = val.trim().toLowerCase();
    document.getElementById('searchClear').classList.toggle('hidden', !activeSearch);
    applyFilters();
  }}, 150);
}}

function clearSearch() {{
  document.getElementById('searchInput').value = '';
  activeSearch = '';
  document.getElementById('searchClear').classList.add('hidden');
  applyFilters();
}}

function filterCat(btn) {{
  hideSchedule();
  document.querySelectorAll('.tab').forEach(b => b.classList.remove('on'));
  btn.classList.add('on');
  activeCat = btn.dataset.cat;
  applyFilters();
  window.scrollTo({{top:0,behavior:'smooth'}});
}}

function filterMood(btn) {{
  const mood = btn.dataset.mood;
  if (activeMood === mood) {{
    clearMood();
    return;
  }}
  document.querySelectorAll('.mood-chip').forEach(c => c.classList.remove('on'));
  btn.classList.add('on');
  activeMood = mood;
  document.getElementById('moodClear').classList.remove('hidden');
  applyFilters();
}}

function clearMood() {{
  activeMood = null;
  document.querySelectorAll('.mood-chip').forEach(c => c.classList.remove('on'));
  document.getElementById('moodClear').classList.add('hidden');
  applyFilters();
}}

// ── Export popup + docx download ──
let _exportTimer = null;
let _exportFile  = null;

function showExportPopup(filename) {{
  _exportFile = filename;
  const overlay   = document.getElementById('exportPopup');
  const btn       = document.getElementById('exportGoBtn');
  const countdown = document.getElementById('exportCountdown');
  const sub       = document.getElementById('exportPopupSub');

  // Reset state
  btn.textContent = "Let's Go ⬇";
  btn.classList.remove('ready');
  sub.innerHTML   = 'Downloading in <span id="exportCountdown">4</span>s…';
  document.getElementById('exportPopupGif').src = 'https://media.giphy.com/media/VK2JbAI71xTxlSVNNu/giphy.gif';
  overlay.classList.remove('hidden');

  let secs = 4;
  document.getElementById('exportCountdown').textContent = secs;

  clearInterval(_exportTimer);
  _exportTimer = setInterval(() => {{
    secs--;
    const el = document.getElementById('exportCountdown');
    if (el) el.textContent = secs;
    if (secs <= 0) {{
      clearInterval(_exportTimer);
      // Can't auto-download on iOS — switch to "tap" state instead
      sub.textContent = 'Ready! Tap the button to save.';
      btn.textContent = '⬇ Download Now';
      btn.classList.add('ready');
    }}
  }}, 1000);

  // window.open as a direct user-gesture call — works on iOS Safari
  btn.onclick = () => {{
    clearInterval(_exportTimer);
    // Swap to Pikachu, then download after 1s
    document.getElementById('exportPopupGif').src = 'https://media.giphy.com/media/ux2EQfCsCm3hJ0dZGv/giphy.gif';
    sub.textContent = '🎉 Downloading…';
    btn.disabled = true;
    const file = _exportFile;
    _exportFile = null;
    setTimeout(() => {{
      document.getElementById('exportPopup').classList.add('hidden');
      btn.disabled = false;
      if (file) window.open(file, '_blank');
    }}, 3000);
  }};
}}

document.addEventListener('keydown', e => {{ if (e.key === 'Escape') {{ closeModal(); closeDrawer(); closeMoodsInfo(); closeStaffGuide(); }} }});

// ── Moods info modal ──
const MOOD_INFO = [
  {{ key:'wind-down',      icon:'😴', name:'Wind Down',
     science:'Myrcene binds GABA-A receptors causing sedation and muscle relaxation (Russo 2011). Linalool elevates adenosine and suppresses glutamate excitability. Together they create the classic "couch-lock" body stone ideal for sleep or unwinding.',
     terps:['Myrcene','Linalool','Caryophyllene'] }},
  {{ key:'anxiety-relief', icon:'🧘', name:'Anxiety Relief',
     science:'Linalool raises GABA and lowers cortisol — the same mechanism as benzodiazepines but milder. Caryophyllene selectively activates CB2 (not CB1) reducing neuroinflammation. Limonene targets 5-HT1A serotonin receptors. Together this is the Kamal 2018 anxiolytic chemotype.',
     terps:['Linalool','Caryophyllene','Limonene'] }},
  {{ key:'lift-up',        icon:'⬆', name:'Lift Up',
     science:'Limonene elevates serotonin and dopamine within 10 minutes of inhalation (Komori 1995, Neuroimmunomodulation). Terpinolene adds cerebral, energizing uplift. Ocimene and Valencene contribute citrus energy without sedation.',
     terps:['Limonene','Terpinolene','Ocimene','Valencene'] }},
  {{ key:'get-creative',   icon:'🎨', name:'Get Creative',
     science:'α-Pinene inhibits acetylcholinesterase (AChE), the enzyme that breaks down acetylcholine — sharpening memory and focus (Miyazawa & Yamafuji 2005). This counters THC-induced short-term memory impairment and drives cerebral, focused creativity.',
     terps:['Pinene','B Pinene','Terpinolene','Limonene'] }},
  {{ key:'get-social',     icon:'😄', name:'Get Social',
     science:'Terpinolene produces euphoria and lowers social inhibitions via serotonergic pathways. Limonene raises dopamine (the reward/motivation neurotransmitter). Combined they produce the giggly, talkative, social-butterfly high.',
     terps:['Limonene','Terpinolene','Ocimene'] }},
  {{ key:'pain-body',      icon:'💆', name:'Pain & Body',
     science:'Caryophyllene is the only terpene that directly binds cannabinoid receptors (CB2 agonist, Gertsch 2008 PNAS) — reducing neuroinflammation and pain signaling. Myrcene is analgesic. Humulene suppresses prostaglandins (same mechanism as ibuprofen).',
     terps:['Caryophyllene','Myrcene','Humulene','Linalool'] }},
  {{ key:'just-happy',     icon:'✨', name:'Just Happy',
     science:'Limonene + Linalool create a balanced euphoria without overstimulation. Terpinolene adds warmth and a gentle creative edge. This is the classic "feel-good" terpene trio — mood-elevating but grounded.',
     terps:['Limonene','Linalool','Terpinolene'] }},
  {{ key:'aphrodisiac',    icon:'🌹', name:'Aphrodisiac',
     science:'Cannabis aphrodisiac use is documented across 3,000 years in India, Persia, and China (Russo 2011). Mechanistically: Limonene raises dopamine (desire), Linalool eliminates anxiety (the #1 arousal blocker), Geraniol (rose terpene) has historic romance associations, Caryophyllene (CB2) may enhance blood flow and tactile sensitivity, Terpinolene lowers inhibitions.',
     terps:['Limonene','Linalool','Geraniol','Caryophyllene','Terpinolene','Ocimene'] }},
];

function openMoodsInfo() {{
  const container = document.getElementById('moodsInfoCards');
  container.innerHTML = MOOD_INFO.map(m => `
    <div class="mood-card">
      <div class="mood-card-head">
        <span class="mood-card-icon">${{m.icon}}</span>
        <span class="mood-card-name">${{m.name}}</span>
      </div>
      <div class="mood-card-science">${{m.science}}</div>
      <div class="mood-card-terps">${{m.terps.map(t=>`<span class="mood-card-terp">${{t}}</span>`).join('')}}</div>
    </div>`).join('');
  document.getElementById('moodsInfoModal').classList.remove('hidden');
  document.body.style.overflow = 'hidden';
}}

function closeMoodsInfo() {{
  document.getElementById('moodsInfoModal').classList.add('hidden');
  document.body.style.overflow = '';
}}

// ── Staff guide modal ──
function openStaffGuide() {{
  const container = document.getElementById('staffGuideContent');
  container.innerHTML = `
    <div class="sg-guide-section">
      <div class="sg-guide-section-title">📱 Using the Menu</div>
      <div class="sg-guide-card">
        <div class="sg-guide-card-head"><span class="sg-guide-card-icon">🗂️</span><span class="sg-guide-card-name">Category Tabs</span></div>
        <div class="sg-guide-card-body">Tap <strong>Flower · Pre-Roll · Vapes · Edibles</strong> at the top to filter by type. Products added in the last 3 days appear in the <strong>✨ New</strong> section automatically.</div>
      </div>
      <div class="sg-guide-card">
        <div class="sg-guide-card-head"><span class="sg-guide-card-icon">🔍</span><span class="sg-guide-card-name">Search</span></div>
        <div class="sg-guide-card-body">The search bar scans everything — product names, terpenes, aroma descriptions, lineage, and therapeutic uses. Try:<br>
          <span class="sg-guide-tag">anxiety</span><span class="sg-guide-tag">sleep</span><span class="sg-guide-tag">Myrcene</span><span class="sg-guide-tag">citrus</span><span class="sg-guide-tag">Cookies</span><span class="sg-guide-tag">PTSD</span>
        </div>
      </div>
      <div class="sg-guide-card">
        <div class="sg-guide-card-head"><span class="sg-guide-card-icon">🎯</span><span class="sg-guide-card-name">Mood Filter</span></div>
        <div class="sg-guide-card-body">Tap any mood chip to score and sort every card 1–10 for that vibe. Border color shows match strength at a glance:<br><br>
          <strong style="color:#16a34a">🟢 Green border</strong> — strong match (7+)<br>
          <strong style="color:#ca8a04">🟡 Amber border</strong> — decent match (4–6)<br>
          <strong style="color:#6b7280">⚫ Gray border</strong> — weak match (1–3)<br><br>
          Tap <strong>ℹ️ How it works</strong> next to the mood chips for the full science breakdown.
        </div>
      </div>
      <div class="sg-guide-card">
        <div class="sg-guide-card-head"><span class="sg-guide-card-icon">🪟</span><span class="sg-guide-card-name">Strain Guide (Tap a Card)</span></div>
        <div class="sg-guide-card-body">Tap any product card to open its full strain profile — lineage, therapeutic uses, aroma, terpenes, and general notes. Hit <strong>＋ Add to Profile</strong> to collect strains for a custom export.</div>
      </div>
    </div>

    <div class="sg-guide-section">
      <div class="sg-guide-section-title">🗂️ Where the Info Comes From</div>
      <div class="sg-guide-card">
        <div class="sg-guide-card-head"><span class="sg-guide-card-icon">💳</span><span class="sg-guide-card-name">Sweed POS — Product Data & Terpenes</span></div>
        <div class="sg-guide-card-body">Name, brand, price, THC/CBD, weight, category, and <strong>terpenes</strong> all come directly from the Sweed POS system. Terpenes in particular come from the <strong>COA (Certificate of Analysis)</strong> each brand submits when they deliver product.<br><br>
          <div class="sg-guide-note">⚠️ If terpenes look incomplete, the fix is in Sweed — update the COA data there and it'll reflect here on the next daily update (4:30 PM CST).</div>
        </div>
      </div>
      <div class="sg-guide-card">
        <div class="sg-guide-card-head"><span class="sg-guide-card-icon">🤖</span><span class="sg-guide-card-name">Claude AI — Strain Profiles</span></div>
        <div class="sg-guide-card-body">When a new strain appears, Claude AI generates its profile using published strain databases and cannabis genetics research. It fills in:<br><br>
          <span class="sg-guide-tag">Lineage</span><span class="sg-guide-tag">Therapeutic uses</span><span class="sg-guide-tag">Side effects</span><span class="sg-guide-tag">Aroma</span><span class="sg-guide-tag">Breeder notes</span><br><br>
          Profiles are generated <em>once per strain</em> and stored. They won't change unless manually updated.
        </div>
      </div>
      <div class="sg-guide-card">
        <div class="sg-guide-card-head"><span class="sg-guide-card-icon">🔬</span><span class="sg-guide-card-name">Published Research — Mood Scoring</span></div>
        <div class="sg-guide-card-body">Mood scores aren't invented — they're grounded in peer-reviewed pharmacology papers:<br><br>
          <strong>Russo 2011</strong> (Br J Pharmacol) — cannabis terpene synergy and therapeutic applications<br>
          <strong>Kamal et al. 2018</strong> (Front Neurosci) — terpene combinations driving anxiolytic effects<br>
          <strong>Gertsch 2008</strong> (PNAS) — Caryophyllene as the only terpene that activates cannabinoid receptors (CB2)
        </div>
      </div>
    </div>

    <div class="sg-guide-section">
      <div class="sg-guide-section-title">🎯 How Mood Scores Are Calculated</div>
      <div class="sg-guide-card">
        <div class="sg-guide-card-body">Scores are <strong>position-weighted</strong> — terpene order on the COA reflects concentration (highest first). The dominant terpene contributes far more than a trace one:<br><br>
          <table class="sg-guide-table">
            <tr><th>COA Position</th><th>Points</th></tr>
            <tr><td>1st (dominant)</td><td>4.0 pts</td></tr>
            <tr><td>2nd</td><td>2.5 pts</td></tr>
            <tr><td>3rd</td><td>1.5 pts</td></tr>
            <tr><td>4th+ (minor)</td><td>0.75 pts</td></tr>
          </table><br>
          Example — <strong>Pain &amp; Body</strong>: Caryophyllene is the only terpene proven to activate CB2 receptors (pain/inflammation). Listed 1st → likely scores 7–8. Listed 3rd → scores 4–5. Absent → max 5 regardless of other terpenes.<br><br>
          <em>Claude AI provides its own 1–10 rating using this same logic plus its knowledge of each strain. When Claude ratings exist, they take priority over the formula.</em>
        </div>
      </div>
    </div>

    <div class="sg-guide-section">
      <div class="sg-guide-section-title">⬇️ Downloading Strain Guides</div>
      <div class="sg-guide-card">
        <div class="sg-guide-card-body">
          <table class="sg-guide-table">
            <tr><th>Button</th><th>Contents</th></tr>
            <tr><td>✅ Available Now</td><td>Every product currently in stock, sorted Flower → Pre-Roll → Vapes</td></tr>
            <tr><td>📦 Master Cache</td><td>Every strain ever enriched (including past products), same sort order</td></tr>
          </table><br>
          Click either button → popup appears → press <strong>Let's Go</strong> → <code>.docx</code> file downloads. Open in Microsoft Word or Google Docs.<br><br>
          Documents rebuild automatically every day. A new strain added to Sweed today will appear in tomorrow's download.
        </div>
      </div>
      <div class="sg-guide-card">
        <div class="sg-guide-card-head"><span class="sg-guide-card-icon">📋</span><span class="sg-guide-card-name">Build a Custom Profile</span></div>
        <div class="sg-guide-card-body">Tap any product card → press <strong>＋ Add to Profile</strong> → open the <strong>📋 My Profile</strong> button (bottom-right corner) → press <strong>⬇ Download Strain Guide</strong> to save an HTML file you can open in Word. Good for building a handpicked reference sheet for a specific customer or use case.</div>
      </div>
    </div>
    <div style="height:20px"></div>
  `;
  document.getElementById('staffGuideModal').classList.remove('hidden');
  document.body.style.overflow = 'hidden';
}}

function closeStaffGuide() {{
  document.getElementById('staffGuideModal').classList.add('hidden');
  document.body.style.overflow = '';
}}

// ── Dark mode ──
function toggleDark() {{
  const dark = document.body.classList.toggle('dark');
  localStorage.setItem('lc-dark', dark ? '1' : '0');
  document.getElementById('darkToggle').textContent = dark ? '☀️ Light' : '🌙 Dark';
}}
(function() {{
  if (localStorage.getItem('lc-dark') === '1') {{
    document.body.classList.add('dark');
    document.addEventListener('DOMContentLoaded', () => {{
      document.getElementById('darkToggle').textContent = '☀️ Light';
    }});
  }}
}})();
</script>

<!-- ── PIN overlay ── -->
<div class="pin-overlay hidden" id="pinOverlay">
  <div class="pin-box">
    <h2>Staff Access</h2>
    <p>Enter your code to view the schedule</p>
    <input class="pin-input" id="pinInput" type="password" maxlength="4"
           inputmode="numeric" placeholder="••••" autocomplete="off"
           oninput="checkPin(this.value)">
    <div class="pin-error" id="pinError"></div>
  </div>
</div>

<!-- ── Schedule section (injected into main by JS) ── -->
<template id="scheduleTemplate">
  <section id="scheduleSection">
    <div class="sched-header">
      <span class="sched-title">📅 Staff Schedule</span>
      <div class="sched-nav">
        <button class="sched-nav-btn" id="schedPrevBtn" onclick="schedNav(-7)">← Previous</button>
        <span class="sched-week-label" id="schedWeekLabel"></span>
        <button class="sched-nav-btn" id="schedNextBtn" onclick="schedNav(7)">Next →</button>
      </div>
    </div>
    <div class="sched-img-section">
      <button class="sched-img-toggle" onclick="toggleSchedImages(this)">📷 View Original Schedule Images</button>
      <div class="sched-img-wrap hidden" id="schedImages">
        <p class="sched-img-label">May 2026</p>
        <img src="schedule-may.jpg" alt="May schedule" class="sched-img">
        <p class="sched-img-label">June 2026</p>
        <img src="schedule-june.jpg" alt="June schedule" class="sched-img">
      </div>
    </div>
    <div class="sched-filters" id="schedFilters"></div>
    <div id="schedDays"></div>
    <div class="sched-updated" id="schedUpdated"></div>
  </section>
</template>

<script>
const SCHEDULE = {schedule_js};
const SCHED_PIN = '0420';
let schedOffset   = 0;   // days from today (multiples of 7)
let schedPerson   = 'all';
let schedUnlocked = false;

function openScheduleTab(btn) {{
  if (schedUnlocked || sessionStorage.getItem('sched-ok') === '1') {{
    schedUnlocked = true;
    showSchedule(btn);
    return;
  }}
  document.getElementById('pinOverlay').classList.remove('hidden');
  document.getElementById('pinInput').value = '';
  document.getElementById('pinError').textContent = '';
  setTimeout(() => document.getElementById('pinInput').focus(), 80);
  // store btn ref to activate after unlock
  window._schedBtn = btn;
}}

function checkPin(val) {{
  if (val.length < 4) return;
  if (val === SCHED_PIN) {{
    sessionStorage.setItem('sched-ok', '1');
    schedUnlocked = true;
    document.getElementById('pinOverlay').classList.add('hidden');
    showSchedule(window._schedBtn);
  }} else {{
    document.getElementById('pinError').textContent = 'Incorrect code';
    document.getElementById('pinInput').value = '';
  }}
}}

function showSchedule(btn) {{
  // Activate tab
  document.querySelectorAll('.tab').forEach(t => t.classList.remove('on'));
  if (btn) btn.classList.add('on');
  // Hide product content, show schedule
  document.querySelectorAll('.section,.new-arrivals-section,.sold-section,.section-divider,.mood-bar,.legend').forEach(el => {{
    el.style.display = 'none';
  }});
  let sec = document.getElementById('scheduleSection');
  if (!sec) {{
    const tpl = document.getElementById('scheduleTemplate');
    const clone = tpl.content.cloneNode(true);
    document.querySelector('main').appendChild(clone);
    sec = document.getElementById('scheduleSection');
  }}
  sec.classList.add('active');
  renderSchedule();
}}

function hideSchedule() {{
  const sec = document.getElementById('scheduleSection');
  if (sec) sec.classList.remove('active');
  document.querySelectorAll('.section,.new-arrivals-section,.sold-section,.section-divider,.mood-bar,.legend').forEach(el => {{
    el.style.display = '';
  }});
}}

function toggleSchedImages(btn) {{
  const wrap = document.getElementById('schedImages');
  const hidden = wrap.classList.toggle('hidden');
  btn.textContent = hidden ? '📷 View Original Schedule Images' : '📷 Hide Schedule Images';
}}

function schedNav(days) {{
  const newOffset = schedOffset + days;
  // Don't go more than 35 days back or 7 days forward
  if (newOffset < -35 || newOffset > 7) return;
  schedOffset = newOffset;
  renderSchedule();
}}

function setSchedPerson(name, btn) {{
  schedPerson = name;
  document.querySelectorAll('.sched-filter-btn').forEach(b => b.classList.remove('on'));
  btn.classList.add('on');
  renderSchedule();
}}

function renderSchedule() {{
  const shifts = SCHEDULE.shifts || [];
  const today  = new Date();
  today.setHours(0,0,0,0);

  // Week window: schedOffset days from today
  const windowStart = new Date(today);
  windowStart.setDate(windowStart.getDate() + schedOffset);
  const windowEnd = new Date(windowStart);
  windowEnd.setDate(windowEnd.getDate() + 6);

  // Nav label
  const fmt = d => d.toLocaleDateString('en-US', {{month:'short', day:'numeric'}});
  document.getElementById('schedWeekLabel').textContent = fmt(windowStart) + ' – ' + fmt(windowEnd);

  // Nav buttons
  document.getElementById('schedPrevBtn').classList.toggle('disabled', schedOffset <= -35);
  document.getElementById('schedNextBtn').classList.toggle('disabled', schedOffset >= 7);

  // Build person filter chips
  const people = ['all', ...Array.from(new Set(shifts.map(s => s.name))).sort()];
  const filtersEl = document.getElementById('schedFilters');
  if (filtersEl.children.length === 0) {{
    people.forEach(p => {{
      const b = document.createElement('button');
      b.className = 'sched-filter-btn' + (p === schedPerson ? ' on' : '');
      b.textContent = p === 'all' ? '👥 Everyone' : p;
      b.onclick = () => setSchedPerson(p, b);
      filtersEl.appendChild(b);
    }});
  }}

  // Render 7 days
  const daysEl = document.getElementById('schedDays');
  daysEl.innerHTML = '';
  for (let i = 0; i < 7; i++) {{
    const day = new Date(windowStart);
    day.setDate(day.getDate() + i);
    const dayStr = day.toISOString().slice(0,10);
    const isToday = day.getTime() === today.getTime();

    const dayShifts = shifts.filter(s => {{
      if (s.date !== dayStr) return false;
      return schedPerson === 'all' || s.name === schedPerson;
    }});

    const dayEl = document.createElement('div');
    dayEl.className = 'sched-day' + (isToday ? ' today' : '');

    const headEl = document.createElement('div');
    headEl.className = 'sched-day-head';
    headEl.innerHTML = day.toLocaleDateString('en-US', {{weekday:'long', month:'short', day:'numeric'}})
      + (isToday ? ' <span class="sched-today-badge">Today</span>' : '');
    dayEl.appendChild(headEl);

    if (dayShifts.length === 0) {{
      const emp = document.createElement('div');
      emp.className = 'sched-shift';
      emp.style.color = 'var(--muted)';
      emp.style.fontStyle = 'italic';
      emp.style.fontSize = '.78rem';
      emp.textContent = 'No shifts scheduled';
      dayEl.appendChild(emp);
    }} else {{
      const t2m = t => {{ const m=t.match(/(\d+):(\d+)\s*(AM|PM)/i); if(!m)return 0; let h=+m[1],pm=m[3].toUpperCase()==='PM'; if(pm&&h!==12)h+=12; if(!pm&&h===12)h=0; return h*60+(+m[2]); }};
      dayShifts.sort((a,b) => t2m(a.start||'') - t2m(b.start||''));
      dayShifts.forEach(s => {{
        const row = document.createElement('div');
        row.className = 'sched-shift';
        row.innerHTML = `<span class="sched-shift-name">${{s.name}}</span>`
          + `<span class="sched-shift-time">${{s.start}} – ${{s.end}}</span>`;
        dayEl.appendChild(row);
      }});
    }}
    daysEl.appendChild(dayEl);
  }}

  // Last updated
  const upd = document.getElementById('schedUpdated');
  if (SCHEDULE.last_updated) {{
    upd.textContent = 'Schedule last updated: ' + new Date(SCHEDULE.last_updated).toLocaleString('en-US', {{month:'short',day:'numeric',year:'numeric',hour:'numeric',minute:'2-digit'}});
  }}
}}
</script>
</body>
</html>"""

    OUT.write_text(html, encoding="utf-8")
    print(f"Built → {OUT}  ({len(all_p)} products)")

    # Also regenerate both docx exports
    build_docx(all_p, db["products"], strains)


DOCX_CAT_ORDER = ["flower", "pre-roll", "vapes"]
DOCX_CAT_LABELS = {"flower": "FLOWER", "pre-roll": "PRE-ROLL", "vapes": "VAPES"}


def _docx_section_header(doc, title):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{'─' * 20}   {title}   {'─' * 20}")
    run.bold = True
    run.font.size = Pt(13)
    doc.add_paragraph()


def _docx_strain(doc, name, product, enriched):
    strain_type = product.get("strain_type", "")
    p = doc.add_paragraph()
    r1 = p.add_run(name)
    r1.bold = True
    r1.font.size = Pt(16)
    p.add_run("\t").bold = True
    r3 = p.add_run(f"-   {strain_type}")
    r3.bold = False
    r3.font.size = Pt(14)

    fields = [
        ("Lineage",     enriched.get("lineage", "")),
        ("Effects",     ", ".join(product.get("effects") or [])),
        ("Flavors",     ", ".join(product.get("flavors") or [])),
        ("Terpenes",    ", ".join(product.get("terpenes") or [])),
        ("Therapeutic", enriched.get("therapeutic", "")),
        ("Negative",    enriched.get("negative", "")),
        ("Aroma",       enriched.get("aroma", "")),
        ("Misc.",       enriched.get("misc", "")),
    ]
    for label, value in fields:
        if not value:
            continue
        p = doc.add_paragraph()
        rb = p.add_run(f"{label}: ")
        rb.bold = True
        p.add_run(value).bold = False
    doc.add_paragraph()


def build_docx(all_p, products_db, strains):
    docs_dir = Path(__file__).parent / "docs"

    def _write(path, items, title):
        doc = Document()
        doc.core_properties.title = title
        # Title paragraph
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = heading.add_run(title.upper())
        hr.bold = True
        hr.font.size = Pt(18)
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(f"MN Legit Cannabis · South Metro · {datetime.now(CST).strftime('%B %d, %Y')}")

        by_cat = {c: [] for c in DOCX_CAT_ORDER}
        for key, p in items:
            cat = (p.get("category") or "").lower()
            if cat in by_cat:
                by_cat[cat].append((key, p))

        for cat in DOCX_CAT_ORDER:
            entries = by_cat[cat]
            if not entries:
                continue
            _docx_section_header(doc, DOCX_CAT_LABELS[cat])
            for key, p in sorted(entries, key=lambda x: x[1].get("name", "")):
                enriched = strains.get(key, {})
                _docx_strain(doc, p.get("name", key), p, enriched)

        doc.save(path)
        print(f"Built docx → {path}  ({sum(len(v) for v in by_cat.values())} strains)")

    # Available Now — only in-stock scraped products
    _write(docs_dir / "legit-available-guide.docx", all_p, "Available Now")

    # Master Cache — all enriched keys, supplemented by products_db
    master_keys = list({**{k: products_db[k] for k in products_db}, **{}}.keys())
    master_items = []
    for k in strains:
        p = products_db.get(k, {"name": k, "category": "flower"})
        master_items.append((k, p))
    _write(docs_dir / "legit-master-guide.docx", master_items, "Master Strain Cache")


build()
